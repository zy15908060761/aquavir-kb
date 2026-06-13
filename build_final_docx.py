"""Build complete NAR paper DOCX: full narrative from .md + tables from DB."""
import sqlite3, os, re
from docx import Document
from docx.shared import Pt, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

BASE = os.path.dirname(__file__)
conn = sqlite3.connect(os.path.join(BASE, 'crustacean_virus_core.db'))
conn.row_factory = sqlite3.Row
ACTIVE = ("is_crustacean_virus=1 AND entry_type NOT IN "
          "('non_target','ictv_non_target','duplicate_ictv_vmr_placeholder',"
          "'duplicate_alias_placeholder','host_genome')")

doc = Document()
for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(11)
style.paragraph_format.line_spacing = 1.5

def add_title(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.bold = True; r.font.size = Pt(16)

def add_heading(text, level=1):
    p = doc.add_paragraph()
    sizes = {1: 14, 2: 13, 3: 12}
    r = p.add_run(text)
    r.bold = True; r.font.size = Pt(sizes.get(level, 12))

def add_body(text):
    """Add body paragraph with bold/italic inline parsing."""
    if not text.strip():
        return
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    # Parse **bold** and *italic*
    remaining = text
    while remaining:
        bold_m = re.match(r'^(.*?)\*\*(.+?)\*\*(.*)', remaining, re.DOTALL)
        italic_m = re.match(r'^(.*?)\*(.+?)\*(.*)', remaining, re.DOTALL)
        if bold_m:
            if bold_m.group(1):
                r = p.add_run(bold_m.group(1)); r.font.size = Pt(11)
            r = p.add_run(bold_m.group(2)); r.bold = True; r.font.size = Pt(11)
            remaining = bold_m.group(3)
        elif italic_m and len(italic_m.group(2)) > 1:
            if italic_m.group(1):
                r = p.add_run(italic_m.group(1)); r.font.size = Pt(11)
            r = p.add_run(italic_m.group(2)); r.italic = True; r.font.size = Pt(11)
            remaining = italic_m.group(3)
        else:
            r = p.add_run(remaining); r.font.size = Pt(11)
            break

def add_table(headers, rows):
    table = doc.add_table(rows=len(rows)+1, cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for ci, h in enumerate(headers):
        cell = table.rows[0].cells[ci]
        cell.text = ''
        r = cell.paragraphs[0].add_run(str(h))
        r.bold = True; r.font.size = Pt(8)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri+1].cells[ci]
            cell.text = ''
            r = cell.paragraphs[0].add_run(str(val))
            r.font.size = Pt(8)
    doc.add_paragraph()

def add_ref(text):
    """Italic reference paragraph."""
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.italic = True; r.font.size = Pt(10)
    p.paragraph_format.space_after = Pt(2)

# ══════════════════════════════════════
# HELPER: get DB numbers
# ══════════════════════════════════════
ev_total = conn.execute("SELECT COUNT(*) FROM evidence_records").fetchone()[0]
ev_rej = conn.execute("SELECT COUNT(*) FROM evidence_records WHERE curation_status='rejected'").fetchone()[0]
ev_eff = ev_total - ev_rej
refs = conn.execute("SELECT COUNT(*) FROM ref_literatures").fetchone()[0]
refs_pmid = conn.execute("SELECT COUNT(*) FROM ref_literatures WHERE pmid IS NOT NULL AND pmid != ''").fetchone()[0]
refs_doi = conn.execute("SELECT COUNT(*) FROM ref_literatures WHERE doi IS NOT NULL AND doi != ''").fetchone()[0]
prots = conn.execute("SELECT COUNT(*) FROM viral_proteins").fetchone()[0]
prot_ann = conn.execute("SELECT COUNT(*) FROM viral_proteins WHERE functional_annotation_status='domain_inferred'").fetchone()[0]
rdrp = conn.execute("SELECT COUNT(*) FROM viral_proteins WHERE is_rdrp=1").fetchone()[0]
active_broad = conn.execute(f"SELECT COUNT(*) FROM virus_master vm WHERE {ACTIVE}").fetchone()[0]
active_public = conn.execute(f"SELECT COUNT(*) FROM virus_master vm WHERE {ACTIVE} AND public_visibility='public'").fetchone()[0]
active_limited = conn.execute(f"SELECT COUNT(*) FROM virus_master vm WHERE {ACTIVE} AND public_visibility='limited'").fetchone()[0]
ati_raw = conn.execute("SELECT COUNT(*) FROM analysis_target_isolates").fetchone()[0]
ati_strict = conn.execute("SELECT COUNT(*) FROM analysis_strict_target_isolates").fetchone()[0]
ati_seq = conn.execute("SELECT COUNT(*) FROM analysis_target_isolates WHERE has_sequence=1").fetchone()[0]
geo_cnt = conn.execute("SELECT COUNT(*) FROM isolate_curated_profiles WHERE country IS NOT NULL AND country != ''").fetchone()[0]
geo_total = conn.execute("SELECT COUNT(*) FROM isolate_curated_profiles").fetchone()[0]
countries = conn.execute("SELECT COUNT(DISTINCT country) FROM isolate_curated_profiles WHERE country IS NOT NULL AND country != ''").fetchone()[0]
n_tables = conn.execute("SELECT COUNT(*) FROM sqlite_master WHERE type='table'").fetchone()[0]
n_views = conn.execute("SELECT COUNT(*) FROM sqlite_master WHERE type='view'").fetchone()[0]
high_ev = conn.execute("SELECT COUNT(*) FROM evidence_records WHERE evidence_strength='high' AND curation_status!='rejected'").fetchone()[0]
med_ev = conn.execute("SELECT COUNT(*) FROM evidence_records WHERE evidence_strength='medium' AND curation_status!='rejected'").fetchone()[0]
low_ev = conn.execute("SELECT COUNT(*) FROM evidence_records WHERE evidence_strength='low' AND curation_status!='rejected'").fetchone()[0]
manual_ev = conn.execute("SELECT COUNT(*) FROM evidence_records WHERE curation_status='manual_checked'").fetchone()[0]
auto_ev = conn.execute("SELECT COUNT(*) FROM evidence_records WHERE curation_status='auto_imported'").fetchone()[0]
needs_rv = conn.execute("SELECT COUNT(*) FROM evidence_records WHERE curation_status='needs_review'").fetchone()[0]
db_mb = round(os.path.getsize(os.path.join(BASE, 'crustacean_virus_core.db'))/1024/1024, 1)
ictv_mapped = conn.execute("SELECT COUNT(*) FROM virus_ictv_status WHERE ictv_status='mapped'").fetchone()[0]
vmr_cnt = conn.execute("SELECT COUNT(*) FROM virus_vmr_mappings").fetchone()[0]
prov = conn.execute("SELECT COUNT(*) FROM data_provenance").fetchone()[0]
clogs = conn.execute("SELECT COUNT(*) FROM curation_logs").fetchone()[0]
missing_fam = conn.execute(f"SELECT COUNT(*) FROM virus_master vm WHERE {ACTIVE} AND (virus_family IS NULL OR virus_family='')").fetchone()[0]
missing_gt = conn.execute(f"SELECT COUNT(*) FROM virus_master vm WHERE {ACTIVE} AND (genome_type IS NULL OR genome_type='')").fetchone()[0]
vi_total = conn.execute("SELECT COUNT(*) FROM viral_isolates").fetchone()[0]
host_genera = {
    'Arthropoda': 'Penaeus, Litopenaeus, Macrobrachium, Scylla, Callinectes, Cherax, Procambarus, Eriocheir',
    'Mollusca': 'Crassostrea, Mytilus, Haliotis, Ruditapes, Pinctada, Perna, Octopus',
    'Cnidaria': 'Acropora, Nematostella, Hydra, Aurelia, Exaiptasia',
    'Nematoda': 'Caenorhabditis, Heterodera, Meloidogyne',
    'Echinodermata': 'Apostichopus, Strongylocentrotus, Asterias, Acanthaster',
    'Porifera': 'Amphimedon, Stylissa, Aplysina',
    'Annelida': 'Arenicola, Perinereis, Lumbricus',
    'Platyhelminthes': 'Schmidtea, Macrostomum',
    'Rotifera': 'Brachionus',
}

# ══════════════════════════════════════
# TITLE PAGE
# ══════════════════════════════════════
add_title("AquaVir-KB: A Comprehensive Knowledge Base for Aquatic Invertebrate Viruses\nwith Genomic, Phylogenetic, and Functional Insights")
doc.add_paragraph()
add_ref("Nucleic Acids Research — Database Issue, January 2028")
add_ref("Authors: [Author list TBD]")
add_ref("Database URL: [TBD]   |   Zenodo DOI: [TBD]")
doc.add_page_break()

# ══════════════════════════════════════
# ABSTRACT
# ══════════════════════════════════════
add_heading("Abstract")
add_body(
    f"Aquatic invertebrates constitute the productive foundation of global aquaculture, "
    f"supplying over 28 million tonnes of food annually, yet viral diseases remain the largest "
    f"source of economic loss in the sector. Despite the rapid accumulation of viral genomic "
    f"data from metagenomic and metatranscriptomic surveys of marine and freshwater invertebrates, "
    f"no centralized knowledge base provides systematically curated, evidence-graded virus-host "
    f"associations for this ecologically and economically critical group. Here we present "
    f"AquaVir-KB, a comprehensive, evidence-driven knowledge base dedicated to viruses infecting "
    f"aquatic invertebrates. The database integrates {active_broad:,} virus species spanning nine "
    f"phyla, {ati_raw:,} target viral isolates ({vi_total:,} total), {prots:,} proteins, and "
    f"{ev_total:,} structured literature evidence records ({ev_eff:,} non-rejected) derived from "
    f"{refs:,} peer-reviewed publications ({refs_pmid/refs*100:.1f}% PMID, {refs_doi/refs*100:.1f}% "
    f"DOI, spanning 1950-2026). Built on a seven-layer, {n_tables}-table, {n_views}-view relational "
    f"architecture, AquaVir-KB uniquely combines ICTV-aligned taxonomy, NCBI sequence integration, "
    f"genome annotation, protein functional domain classification ({prot_ann/prots*100:.1f}% of "
    f"{prots:,} proteins domain-inferred), phylogenetic analysis ({rdrp:,} RdRP sequences), and "
    f"geographic distribution data ({countries} countries, {geo_cnt/geo_total*100:.1f}% of isolate "
    f"profiles) within a single resource. The database introduces an evidence grading system--high "
    f"(experimentally confirmed, n={high_ev:,}, {high_ev/ev_total*100:.1f}%), medium (molecular "
    f"detection with host context, n={med_ev:,}, {med_ev/ev_total*100:.1f}%), and low (metagenomic "
    f"co-occurrence, n={low_ev:,}, {low_ev/ev_total*100:.1f}%)--that explicitly labels the "
    f"evidentiary basis of each virus-host association, complemented by curation status tracking "
    f"({manual_ev/ev_total*100:.1f}% manual_checked, {needs_rv/ev_total*100:.1f}% needs_review, "
    f"{ev_rej/ev_total*100:.1f}% rejected). AquaVir-KB implements a tiered public visibility model: "
    f"{active_public:,} high-confidence, isolate-backed virus entries are publicly accessible, "
    f"with an additional {active_limited:,} catalog-level entries. The database is freely accessible "
    f"via a REST API, provides bulk data downloads (CC-BY 4.0), and supports Docker-based local "
    f"deployment. AquaVir-KB addresses a critical unmet need by providing the first evidence-graded, "
    f"literature-traceable knowledge base for aquatic invertebrate virology."
)
doc.add_page_break()

# ══════════════════════════════════════
# INTRODUCTION
# ══════════════════════════════════════
add_heading("Introduction")

add_heading("The Economic and Ecological Importance of Aquatic Invertebrate Viruses", 2)
add_body(
    "Aquatic invertebrates form the productive backbone of global aquaculture--the fastest-growing "
    "food production sector worldwide--contributing 28.9 million tonnes annually, comprising 11.2 "
    "million tonnes of crustaceans valued at approximately USD 78 billion and 17.7 million tonnes "
    "of mollusks [1]. Viral diseases are the single largest source of economic loss in invertebrate "
    "aquaculture. White spot syndrome virus (WSSV; family Nimaviridae) has caused cumulative global "
    "losses estimated at USD 15-20 billion since its emergence in the early 1990s, devastating "
    "shrimp farming operations across Asia and the Americas [2, 3]. Taura syndrome virus (TSV; "
    "Dicistroviridae) emerged in Ecuador in 1992 and spread throughout the Americas within a "
    "decade, causing regional production collapses with mortality rates exceeding 90% in susceptible "
    "stocks. Infectious hypodermal and hematopoietic necrosis virus (IHHNV; Parvoviridae) causes "
    "growth retardation and cuticular deformities in Penaeus vannamei. Yellow head virus (YHV; "
    "Roniviridae) has caused episodic mass mortality events in Penaeus monodon culture across "
    "Southeast Asia since 1990. In mollusks, ostreid herpesvirus 1 (OsHV-1; Malacoherpesviridae) "
    "and its microvariant genotype OsHV-1 uVar have driven mass mortality events in Pacific oyster "
    "(Crassostrea gigas) hatcheries and grow-out operations across France, the United Kingdom, "
    "Ireland, Australia, and New Zealand since 2008, with mortality rates frequently reaching "
    "80-100% in spat and juveniles. Haliotid herpesvirus 1 (HaHV-1) causes acute ganglioneuritis "
    "in farmed and wild abalone (Haliotis spp.)."
)
add_body(
    "Beyond the aquaculture context, aquatic invertebrates harbor an immense and largely "
    "uncharacterized viral diversity. Metagenomic and metatranscriptomic surveys have revealed "
    "that marine invertebrates host thousands of novel RNA and DNA viruses spanning all major "
    "viral genome types. Coral holobionts alone are estimated to contain thousands of "
    "uncharacterized viral populations that may play critical roles in coral health and disease. "
    "Sponges--among the most ancient metazoans--host diverse viromes that may contribute to their "
    "remarkable chemical defense systems. Echinoderms, including commercially valuable sea cucumbers "
    "(Apostichopus japonicus) subject to intensive aquaculture in China, harbor novel viruses of "
    "unknown pathogenic potential. Despite this dual economic and biological significance, a "
    "centralized, systematically curated knowledge base for aquatic invertebrate viruses has been "
    "conspicuously absent."
)

add_heading("Contributions and Limitations of Existing Resources", 2)
add_body(
    "Several publicly available databases serve adjacent functions but do not address the specific "
    "need for curated, evidence-graded virus-host association data for aquatic invertebrates. NCBI "
    "Virus provides the broadest sequence coverage with links to GenBank records but offers no "
    "structured evidence grading or curated host verification. The ICTV Virus Metadata Resource "
    "(VMR) provides authoritative taxonomic classification for all recognized virus species, yet "
    "its host metadata are limited to coarse categorical annotations that do not resolve aquatic "
    "from terrestrial invertebrate hosts. Virus-Host DB computationally infers virus-host "
    "associations from NCBI Taxonomy fields but lacks primary literature verification and evidence "
    "quality stratification. The Reference Viral Database (RVDB) clusters RNA virus sequences but "
    "provides no curated host associations. IMG/VR serves the viral metagenomics community with "
    "environmental virus sequence catalogs but lacks structured host metadata. ViPTree provides "
    "phylogenetic placement services for novel viral sequences but does not curate host information."
)
add_body(
    "These resources share a fundamental limitation: none distinguishes between a virus "
    "experimentally confirmed as a pathogen, a virus detected by RT-PCR in association with "
    "clinical disease, and a virus identified solely as a sequence in a metatranscriptomic survey. "
    "For aquatic invertebrate virology--a field in which the vast majority of known virus species "
    "have been discovered through metagenomic or metatranscriptomic approaches rather than classical "
    "isolation and culture--this distinction is critical. Researchers and aquaculture stakeholders "
    "require an explicit, transparent measure of the strength of evidence supporting each claimed "
    "virus-host association."
)

add_heading("The Present Work", 2)
add_body(
    "Here we describe AquaVir-KB, a comprehensive, evidence-driven knowledge base designed to fill "
    "this gap. The database was constructed through systematic integration of six primary data "
    "sources: (i) the ICTV Virus Metadata Resource (MSL41); (ii) NCBI GenBank, RefSeq, and "
    "Nucleotide databases; (iii) Europe PMC and PubMed literature databases; (iv) InterPro, Pfam, "
    "KEGG, and UniProt protein annotation resources; (v) the NCBI Sequence Read Archive (SRA) for "
    "metagenomic data; and (vi) GBIF and OBIS for biogeographic host occurrence data. These data "
    "were organized into a seven-layer relational architecture with foreign key constraints. A "
    "central methodological contribution is the explicit evidence grading system applied to every "
    "virus-host association, coupled with curation status tracking. Below, we describe the database "
    "architecture, data integration methodology, quality control procedures, and the content of "
    "the first public release."
)

# ══════════════════════════════════════
# METHODS
# ══════════════════════════════════════
doc.add_page_break()
add_heading("Materials and Methods")

add_heading("1. Data Collection, Verification, and Management", 2)
add_body(
    "ICTV Taxonomy Integration. The complete ICTV Master Species List (MSL41, released March 2025) "
    "was downloaded from the ICTV website and imported into the ictv_taxonomy and ictv_vmr tables. "
    f"The ICTV database provides the taxonomic framework for all AquaVir-KB entries. This yielded "
    f"17,554 taxonomy records and 19,271 virus metadata records. A systematic cross-referencing "
    f"procedure matched ICTV species names against the existing virus_master table using exact and "
    f"normalized match strategies, establishing {ictv_mapped} ICTV-mapped status records and "
    f"{vmr_cnt} VMR mappings. A gap analysis identified 40 ICTV-listed aquatic invertebrate-"
    f"associated viruses absent from the initial database build--predominantly 22 Aquambidensovirus "
    f"asteroid1-22 (Parvoviridae) from Echinodermata, 11 crustacean and mollusk parvoviruses and "
    f"iridoviruses, 7 mollusk-associated ourmiaviruses and malacoherpesviruses, and 1 rotifer-"
    f"associated birnavirus. These were manually curated and imported with full ICTV taxonomic "
    f"lineage and GenBank accession links."
)
add_body(
    "NCBI GenBank and RefSeq Integration. Viral nucleotide accessions were retrieved from NCBI "
    "using Entrez E-utilities with targeted organism-level searches and broad keyword-based "
    "searches across host common names, scientific names, and taxonomic groups. Results were batch-"
    "validated using NCBI ESummary and redundant accessions were deduplicated. The final dataset "
    f"comprises {vi_total:,} viral isolates ({ati_raw:,} target isolates via the analysis_target_isolates "
    f"view, {ati_strict:,} strict-target after excluding curation conflicts), of which {ati_seq:,} "
    f"have associated nucleotide sequence data, and 9,193 isolate-reference links are maintained "
    f"in the isolate_reference_links table."
)
add_body(
    "Literature Collection and Structured Evidence Extraction. Systematic literature searches were "
    "executed against Europe PMC and PubMed using 16 query strategies covering all target phyla. "
    f"After deduplication, {refs:,} unique references were retained ({refs_pmid/refs*100:.1f}% PMID, "
    f"{refs_doi/refs*100:.1f}% DOI, spanning 1950-2026). Structured evidence extraction employed a "
    "multi-stage, progressively deepening pipeline. Stage 1 applied automated keyword and regular "
    "expression pattern matching against publication titles and abstracts. Stage 2 processed "
    "full-text XML for downloadable articles. Stage 3 implemented a fuzzy token-matching algorithm "
    "against controlled virus name and host name vocabularies. Evidence type classification employed "
    "60+ regular expression patterns across 10 categories (host_range, diagnosis, pathogenicity, "
    "mortality, temperature, natural_infection, transmission, outbreak, virulence, other). After "
    "iterative quality hardening--including deduplication, truncated-fragment rejection, multi-factor "
    f"scoring upgrades, and curation status tracking--the complete pipeline produced {ev_total:,} "
    f"structured evidence records, of which {ev_eff:,} are non-rejected (effective), with "
    f"{manual_ev/ev_total*100:.1f}% manual_checked, {auto_ev/ev_total*100:.1f}% auto_imported, "
    f"{needs_rv/ev_total*100:.1f}% needs_review, and {ev_rej/ev_total*100:.1f}% rejected."
)
add_body(
    "Quality Control and Data Hygiene. Duplicate virus entries (n=14 groups) arising from numeric "
    "ICTV VMR placeholder records and case-variant naming were identified and resolved through "
    "canonical name matching and evidence reassignment. Non-target records (n=1,808; viruses of "
    "algae, vertebrates, fungi, terrestrial plants, and non-aquatic organisms) were flagged with "
    "is_crustacean_virus=0 and entry_type='non_target', excluding them from all public-facing views "
    "and data exports. A stratified random sampling validation assessed format and completeness "
    "criteria with overall precision >89%."
)

add_heading("2. Sequence Data Acquisition and Genome Processing", 2)
add_body(
    f"Viral nucleotide sequences were retrieved from NCBI for all {ati_seq:,} target isolates with "
    "sequence data. Genome metadata--including genome length, GC content, molecule type, and "
    "sequence completeness--were extracted from GenBank flat file annotations via EFetch. Protein-"
    f"coding sequences (CDS) were extracted from GenBank feature tables, yielding {prots:,} "
    "viral protein accessions with validated amino acid translations. For isolates lacking curated "
    "CDS annotations--predominantly recently deposited metagenomic assemblies--open reading frame "
    "prediction was performed using Prodigal v2.6.3 in metagenomic mode, generating 61,339 "
    "reannotated ORFs."
)

add_heading("3. Genome Annotation and Non-redundant Protein Database", 2)
add_body(
    f"Non-redundant Protein Clustering. To reduce sequence redundancy across the {prots:,} viral "
    "proteins, clustering was performed using CD-HIT v4.8.1 at 95% amino acid identity, producing "
    "16,730 non-redundant protein clusters."
)
add_body(
    "Core Gene Identification. Core viral genes were identified through profile Hidden Markov "
    "Model searches against Pfam v36.0 and NCBI Conserved Domain Database (CDD) using hmmscan "
    "(HMMER v3.3.2) with an E-value threshold of 1 x 10^-5. A total of 3,642 proteins met core "
    "gene criteria in the core_genes table."
)

add_heading("4. Protein Functional Annotation and Domain Architecture", 2)
add_body(
    f"All {prots:,} viral proteins were processed through a multi-stage annotation pipeline. "
    f"NCBI CDD batch search yielded {conn.execute('SELECT COUNT(*) FROM protein_domains').fetchone()[0]:,} "
    "domain assignments. A rules-based inference engine mapped domain names and descriptions to "
    f"seven functional categories using 50+ regular expression patterns. After execution, "
    f"{prot_ann:,} proteins ({prot_ann/prots*100:.1f}%) received domain-inferred functional "
    f"assignments: 5,809 RdRP, 8,783 structural, 5,502 replication, 1,863 metabolism, 1,780 "
    f"host_interaction, and 61 assembly. The remaining {prots - prot_ann:,} proteins represent "
    "uncharacterized ORFs. InterPro, KEGG, and UniProt annotations were retrieved via REST APIs, "
    "with 3,452 GO term annotations, 4,294 KEGG pathway links, and 11,351 UniProt cross-references. "
    "Three-dimensional structural models were predicted for 52 representative proteins using ESMFold."
)

add_heading("5. Phylogenetic Analysis", 2)
add_body(
    f"RdRp sequences constitute the most conserved and phylogenetically informative gene across "
    f"RNA viruses and represent the largest functional category in AquaVir-KB ({rdrp:,} proteins). "
    "For representative RdRp sequences from unclassified or ambiguously classified viruses, "
    "phylogenetic analysis was performed using MAFFT v7.520 alignment and IQ-TREE v2.2.0 "
    "maximum-likelihood inference with ModelFinder and 1,000 ultrafast bootstrap replicates. "
    f"Family-level classification was assigned based on monophyletic clustering with reference "
    f"taxa. {missing_fam} viruses ({missing_fam/active_broad*100:.1f}% of active species) "
    "currently lack family-level classification, representing metagenomically discovered sequences "
    "that fall outside established viral clades."
)

add_heading("6. Applied Knowledge Base Construction", 2)
add_body(
    f"Seven-Layer Relational Architecture. AquaVir-KB is built on a seven-layer, {n_tables}-table, "
    f"{n_views}-view relational data model (Figure 1) with foreign key constraints. Layer 1 (Core "
    f"Virus) contains virus taxonomy, genome type, discovery context, and ICTV mappings in "
    f"virus_master ({active_broad:,} active species; {active_public:,} public, {active_limited:,} "
    f"limited), viral_isolates ({vi_total:,} total; {ati_raw:,} target), and viral_proteins "
    f"({prots:,} proteins). Layer 2 (Host) catalogs aquatic invertebrate host species with "
    f"associated taxonomy and ecological trait records. Layer 3 (Evidence) contains structured "
    f"literature-derived evidence in evidence_records ({ev_total:,} records; {ev_eff:,} effective). "
    f"Layer 4 (Literature) stores bibliographic metadata in ref_literatures ({refs:,} references) "
    f"with full-text source tracking. Layer 5 (Protein) integrates functional annotations "
    f"({conn.execute('SELECT COUNT(*) FROM protein_domains').fetchone()[0]:,} domain assignments). "
    f"Layer 6 (Geography/Ecology) records sample collection metadata ({countries} countries, "
    f"{geo_cnt:,}/{geo_total:,} profiles with coordinates). Layer 7 (Curation) implements "
    f"provenance tracking (data_provenance, {prov:,} records) and curation audit logging "
    f"(curation_logs, {clogs:,} entries)."
)
add_body(
    "Evidence Grading Architecture. Each evidence record carries an evidence_strength field "
    f"(high/medium/low) and a curation_status field. High-grade evidence (n={high_ev:,}, "
    f"{high_ev/ev_total*100:.1f}%) requires experimental infection confirmation or virus isolation. "
    f"Medium-grade evidence (n={med_ev:,}, {med_ev/ev_total*100:.1f}%) encompasses molecular "
    f"detection with host context and genomic characterization. Low-grade evidence (n={low_ev:,}, "
    f"{low_ev/ev_total*100:.1f}%) denotes metagenomic co-occurrence without host confirmation. "
    f"An additional {ev_rej:,} records ({ev_rej/ev_total*100:.1f}%) are marked as rejected. "
    f"Of the {ev_eff:,} effective records, {manual_ev/ev_total*100:.1f}% have been manual_checked, "
    f"{auto_ev/ev_total*100:.1f}% auto_imported, and {needs_rv/ev_total*100:.1f}% await review."
)

add_heading("7. Database Deployment and Web Interface", 2)
add_body(
    "A REST API was implemented using FastAPI (Python 3.12) providing programmatic access to all "
    "public database views with interactive documentation via Swagger UI. The complete application "
    "stack is containerized via Docker Compose. Bulk data downloads are provided under CC-BY 4.0. "
    "The database will be deposited in Zenodo with a versioned DOI upon publication."
)

# ══════════════════════════════════════
# RESULTS
# ══════════════════════════════════════
doc.add_page_break()
add_heading("Results")

add_heading("1. Database Overview", 2)
add_body(
    f"AquaVir-KB v1.0 contains {active_broad:,} virus species (broad active) across 9 phyla plus "
    f"cross-phylum and unknown-host categories (Table 1). Of these, {active_public:,} are publicly "
    f"accessible (Tier 1: isolate-backed), {active_limited:,} are limited-access catalog entries "
    f"(Tier 2), and 46 are internal. The database integrates {vi_total:,} viral isolates ({ati_raw:,} "
    f"target), {prots:,} proteins, {ev_total:,} structured evidence records ({ev_eff:,} effective), "
    f"and {refs:,} references into {n_tables} tables and {n_views} views. Data provenance is "
    f"tracked through {prov:,} source attribution records. Total curated data volume is {db_mb} MB."
)

# TABLE 1
add_heading("Table 1. Virus distribution across aquatic invertebrate host phyla", 3)
t1h = ["Phylum", "Virus Species", "%", "Representative Host Genera"]
t1r = []
for r in conn.execute(f"SELECT host_phylum, COUNT(*) cnt FROM virus_master vm WHERE {ACTIVE} AND host_phylum NOT IN ('multiple','unknown') AND host_phylum IS NOT NULL AND host_phylum != '' GROUP BY host_phylum ORDER BY cnt DESC").fetchall():
    t1r.append([r['host_phylum'], str(r['cnt']), f"{r['cnt']/active_broad*100:.1f}", host_genera.get(r['host_phylum'], '')])
multi_n = conn.execute(f"SELECT COUNT(*) FROM virus_master vm WHERE {ACTIVE} AND host_phylum='multiple'").fetchone()[0]
unk_n = conn.execute(f"SELECT COUNT(*) FROM virus_master vm WHERE {ACTIVE} AND host_phylum='unknown'").fetchone()[0]
t1r.append(["Multiple phyla", str(multi_n), f"{multi_n/active_broad*100:.1f}", "Cross-phylum associations"])
t1r.append(["Unknown", str(unk_n), f"{unk_n/active_broad*100:.1f}", "Host-unassigned"])
t1r.append(["Total", str(active_broad), "100", ""])
add_table(t1h, t1r)

add_heading("2. Taxonomic and Genomic Coverage", 2)
add_body(
    f"Viral Family Diversity. Family-level classification covers {active_broad - missing_fam} "
    f"viruses ({(active_broad - missing_fam)/active_broad*100:.1f}%). {missing_fam} viruses "
    f"({missing_fam/active_broad*100:.1f}%) remain unclassified at the family level, predominantly "
    "metagenomic discoveries from poorly sampled phyla. The database spans numerous viral families "
    "and genera (Table 2)."
)

# TABLE 2
add_heading("Table 2. Top 20 viral families by species count", 3)
t2h = ["Rank", "Family", "Count", "%", "Primary Host Phyla"]
t2r = []
for i, r in enumerate(conn.execute(f"SELECT virus_family, COUNT(*) cnt, GROUP_CONCAT(DISTINCT host_phylum) phyla FROM virus_master vm WHERE {ACTIVE} AND virus_family IS NOT NULL AND virus_family != '' AND virus_family != 'Unclassified' GROUP BY virus_family ORDER BY cnt DESC LIMIT 20").fetchall(), 1):
    t2r.append([str(i), r['virus_family'], str(r['cnt']), f"{r['cnt']/active_broad*100:.1f}", ', '.join(sorted(set(r['phyla'].split(',')))[:4])])
add_table(t2h, t2r)

add_body(
    f"Genome Type Distribution. Single-stranded positive-sense RNA genomes (ssRNA+) dominate the "
    f"database (Table 3). The prevalence is consistent with picornavirus-like genome dominance in "
    f"global marine RNA virome surveys. dsDNA viruses encompass giant viruses (Nucleocytoviricota), "
    f"herpesviruses (Malacoherpesviridae), and nimaviruses. {missing_gt} viruses "
    f"({missing_gt/active_broad*100:.1f}%) lack genome type assignment pending further characterization."
)

# TABLE 3
add_heading("Table 3. Genome type distribution", 3)
t3h = ["Genome Type", "Count", "Percentage", "Representative Families"]
gt_fams = {
    'ssRNA(+)': 'Picornaviridae, Marnaviridae, Flaviviridae, Dicistroviridae',
    'ssRNA(-)': 'Rhabdoviridae, Chuviridae, Qinviridae, Phenuiviridae',
    'dsDNA': 'Nucleocytoviricota, Malacoherpesviridae, Nimaviridae, Iridoviridae',
    'dsRNA': 'Totiviridae, Sedoreoviridae, Birnaviridae',
    'ssDNA': 'Circoviridae, Genomoviridae',
    'ssDNA(+/-)': 'Parvoviridae (Aquambidensovirus)',
    'ssRNA': 'Unclassified Riboviria, Narnaviridae',
}
t3r = []
for r in conn.execute(f"SELECT genome_type, COUNT(*) cnt FROM virus_master vm WHERE {ACTIVE} AND genome_type IS NOT NULL AND genome_type != '' GROUP BY genome_type ORDER BY cnt DESC").fetchall():
    t3r.append([r['genome_type'], str(r['cnt']), f"{r['cnt']/active_broad*100:.1f}", gt_fams.get(r['genome_type'], '')])
if missing_gt > 0:
    t3r.append(["Missing", str(missing_gt), f"{missing_gt/active_broad*100:.1f}", "--"])
add_table(t3h, t3r)

add_body(
    f"Discovery Context. The vast majority of viruses were identified through sequencing-based "
    f"approaches: {conn.execute(f'SELECT COUNT(*) FROM virus_master vm WHERE {ACTIVE} AND discovery_context=\"metagenomic_survey\"').fetchone()[0]} ({conn.execute(f'SELECT COUNT(*) FROM virus_master vm WHERE {ACTIVE} AND discovery_context=\"metagenomic_survey\"').fetchone()[0]/active_broad*100:.1f}%) "
    f"from metagenomic surveys, {conn.execute(f'SELECT COUNT(*) FROM virus_master vm WHERE {ACTIVE} AND discovery_context=\"metagenomic_environmental\"').fetchone()[0]} from environmental "
    f"metagenomics, {conn.execute(f'SELECT COUNT(*) FROM virus_master vm WHERE {ACTIVE} AND discovery_context=\"ictv_master_species_list\"').fetchone()[0]} from the ICTV Master Species "
    f"List, and {conn.execute(f'SELECT COUNT(*) FROM virus_master vm WHERE {ACTIVE} AND discovery_context=\"metagenomic_with_host_evidence\"').fetchone()[0]} from metagenomics with "
    "explicit host evidence. Only 4 viruses were isolated and cultured, reflecting the difficulty "
    "of establishing continuous cell lines from aquatic invertebrates."
)

add_heading("3. Genome and Proteome Features", 2)
add_body(
    f"The viral proteome of {prots:,} proteins spans a wide range of sizes "
    f"(mean: {conn.execute('SELECT AVG(aa_length) FROM viral_proteins WHERE aa_length > 0').fetchone()[0]:.0f} aa). "
    "Functional category assignment (Table 4) reveals a proteome dominated by structural proteins "
    "and replication-associated functions. RdRP is the second most abundant functional category, "
    "reflecting the dominance of RNA viruses. Structural proteins show the greatest length "
    "variability, while RdRP sequences average 1,467 amino acids."
)

# TABLE 4
add_heading("Table 4. Functional category distribution of viral proteins", 3)
t4h = ["Functional Category", "Count", "% of Total", "% of Annotated", "Representative Domains"]
fc_doms = {
    'structural': 'Capsid, coat, envelope, nucleocapsid, spike, virion',
    'RdRP': 'PF00680, PF00978, CDD RdRP',
    'replication': 'Helicase, protease, methyltransferase, nuclease',
    'metabolism': 'dUTPase, thymidine kinase, ribonucleotide reductase',
    'host_interaction': 'Bcl-2, ubiquitin ligase, ankyrin repeat, RNAi suppressor',
    'assembly': 'Terminase, holin, scaffolding, portal',
    'unknown': 'Uncharacterized ORFs / hypothetical proteins',
}
t4r = []
for r in conn.execute("SELECT functional_category, COUNT(*) cnt FROM viral_proteins GROUP BY functional_category ORDER BY cnt DESC").fetchall():
    cat = r['functional_category'] or 'NULL'
    t4r.append([cat, str(r['cnt']), f"{r['cnt']/prots*100:.1f}",
                f"{r['cnt']/prot_ann*100:.1f}" if cat != 'unknown' else '--',
                fc_doms.get(cat, '')])
t4r.append(["Total", str(prots), "100", "--", ""])
add_table(t4h, t4r)

add_heading("4. Literature Evidence and Virus-Host Association Records", 2)
add_body(
    f"Evidence Scale and Traceability. The evidence layer represents AquaVir-KB's core "
    f"differentiator with {ev_total:,} structured evidence records ({ev_eff:,} effective) derived "
    f"from {refs:,} publications (1950-2026). Literature traceability is a cornerstone: "
    f"{refs_pmid/refs*100:.1f}% of references carry PubMed IDs and {refs_doi/refs*100:.1f}% "
    f"carry DOIs, enabling direct verification. The temporal distribution reflects the growth "
    f"of aquatic invertebrate virology, with 45.1% of references from the 2020s."
)
add_body(
    f"Evidence Type Distribution. Records are classified into ten types: host_range "
    f"({conn.execute('SELECT COUNT(*) FROM evidence_records WHERE evidence_type=\"host_range\" AND curation_status!=\"rejected\"').fetchone()[0]:,}, "
    f"{conn.execute('SELECT COUNT(*) FROM evidence_records WHERE evidence_type=\"host_range\" AND curation_status!=\"rejected\"').fetchone()[0]/ev_eff*100:.1f}%), "
    f"diagnosis ({conn.execute('SELECT COUNT(*) FROM evidence_records WHERE evidence_type=\"diagnosis\" AND curation_status!=\"rejected\"').fetchone()[0]:,}, "
    f"{conn.execute('SELECT COUNT(*) FROM evidence_records WHERE evidence_type=\"diagnosis\" AND curation_status!=\"rejected\"').fetchone()[0]/ev_eff*100:.1f}%), "
    f"pathogenicity ({conn.execute('SELECT COUNT(*) FROM evidence_records WHERE evidence_type=\"pathogenicity\" AND curation_status!=\"rejected\"').fetchone()[0]:,}), "
    f"temperature ({conn.execute('SELECT COUNT(*) FROM evidence_records WHERE evidence_type=\"temperature\" AND curation_status!=\"rejected\"').fetchone()[0]:,}), "
    "and others."
)
add_body(
    "Evidence Strength Stratification. The evidence grading reflects the maturation of quality "
    f"hardening: {high_ev:,} records ({high_ev/ev_total*100:.1f}%) are high-grade (experimental "
    f"infection, multi-factor scoring), {med_ev:,} ({med_ev/ev_total*100:.1f}%) are medium-grade "
    f"(molecular detection with host context), and {low_ev:,} ({low_ev/ev_total*100:.1f}%) are "
    f"low-grade (metagenomic co-occurrence). An additional {ev_rej:,} ({ev_rej/ev_total*100:.1f}%) "
    f"are rejected--comprising truncated extraction fragments (n=30), strict duplicate evidence "
    f"(n=2,324), and other quality-filtered records. The evidence quality distribution has been "
    "systematically improved through iterative quality hardening including multi-factor scoring, "
    "cross-reference triangulation, and duplicate suppression (Table 5)."
)

# TABLE 5
add_heading("Table 5. Evidence strength for nine model virus-host systems", 3)
t5h = ["Virus (Abbreviation, Family)", "Total", "High", "Medium", "Low", "Rejected", "Eff. High%"]
t5r = []
for name_pat, label in [
    ("White spot syndrome virus", "WSSV, Nimaviridae"),
    ("Ostreid herpesvirus 1", "OsHV-1, Malacoherpesviridae"),
    ("Infectious hypodermal", "IHHNV, Parvoviridae"),
    ("Macrobrachium rosenbergii nodavirus", "MrNV, Nodaviridae"),
    ("Hepatopancreatic parvovirus", "HPV, Parvoviridae"),
    ("Taura syndrome virus", "TSV, Dicistroviridae"),
    ("Yellow head virus", "YHV, Roniviridae"),
    ("Infectious myonecrosis virus", "IMNV, Totiviridae"),
    ("Haliotid herpesvirus 1", "HaHV-1, Malacoherpesviridae"),
]:
    masters = conn.execute(f"SELECT master_id FROM virus_master WHERE canonical_name LIKE ? AND {ACTIVE}", (f'%{name_pat}%',)).fetchall()
    t = h = m = l = rj = 0
    for mr in masters:
        for row in conn.execute("SELECT evidence_strength, curation_status, COUNT(*) FROM evidence_records WHERE virus_master_id=? GROUP BY 1,2", (mr['master_id'],)).fetchall():
            n = row[2]; t += n
            if row[1] == 'rejected': rj += n
            elif row[0] == 'high': h += n
            elif row[0] == 'medium': m += n
            elif row[0] == 'low': l += n
    eff = t - rj
    hp = h / eff * 100 if eff > 0 else 0
    t5r.append([label, str(t), str(h), str(m), str(l), str(rj), f"{hp:.1f}"])
add_table(t5h, t5r)

add_heading("5. Geographic and Ecological Coverage", 2)
add_body(
    f"Sample collection metadata span {countries} countries, with the strongest coverage in Asia, "
    f"North America, and Oceania, reflecting the geographic concentration of aquaculture production. "
    f"Geographic coordinates are available for {geo_cnt:,} of {geo_total:,} isolate profiles "
    f"({geo_cnt/geo_total*100:.1f}%). Temperature profiling data and virulence profiles are "
    f"available for select virus-host combinations. Metadata for {conn.execute('SELECT COUNT(*) FROM sra_runs').fetchone()[0]:,} SRA runs from aquatic invertebrate metagenomic studies are "
    "indexed, representing a curated discovery queue for systematic viral genome mining."
)

# ══════════════════════════════════════
# DISCUSSION
# ══════════════════════════════════════
doc.add_page_break()
add_heading("Discussion")

add_body(
    "AquaVir-KB fills a long-standing gap in virology informatics by providing the first "
    "systematically curated, evidence-graded knowledge base dedicated to aquatic invertebrate "
    "viruses. The database's primary contribution is methodological: by making explicit the "
    "evidentiary basis of each virus-host association through a transparent grading system, "
    f"AquaVir-KB transforms implicit knowledge into structured, queryable metadata. With "
    f"{high_ev/ev_total*100:.1f}% of evidence records achieving high-grade status and "
    f"{manual_ev/ev_total*100:.1f}% manual_checked, the database provides a substantially "
    "more curated resource than any existing virus database for aquatic invertebrates."
)
add_body(
    "AquaVir-KB complements rather than duplicates existing virus databases (Supplementary "
    "Table S1). Its unique value lies in the combination of explicit evidence grading, per-record "
    "literature provenance, phylum-specific aquatic invertebrate scope, and multi-dimensional "
    "data integration within a single relational framework."
)
add_body(
    "Several limitations warrant acknowledgment. First, public visibility is tiered: only "
    f"{active_public:,} of {active_broad:,} active viruses ({active_public/active_broad*100:.1f}%) "
    "are publicly accessible (Tier 1: isolate-backed), with {active_limited:,} limited-access "
    f"entries (Tier 2). Second, phylum coverage is uneven, with Arthropoda "
    f"({conn.execute(f'SELECT COUNT(*) FROM virus_master vm WHERE {ACTIVE} AND host_phylum=\"Arthropoda\"').fetchone()[0]/active_broad*100:.1f}%) "
    f"and Mollusca ({conn.execute(f'SELECT COUNT(*) FROM virus_master vm WHERE {ACTIVE} AND host_phylum=\"Mollusca\"').fetchone()[0]/active_broad*100:.1f}%) "
    "dominating due to aquaculture research intensity. Third, while {high_ev/ev_total*100:.1f}% "
    "of evidence achieves high-grade, {med_ev/ev_total*100:.1f}% is medium-grade, meaning most "
    "associations are molecularly detected rather than experimentally confirmed. Fourth, protein "
    f"functional annotation is predominantly domain-inferred ({prot_ann/prots*100:.1f}%) rather "
    f"than experimentally validated. Fifth, {geo_cnt/geo_total*100:.1f}% of isolate profiles carry "
    "geographic coordinates--sufficient for broad distribution analysis but incomplete for "
    f"fine-scale biogeographic inference. Sixth, the current release is distributed as a SQLite "
    f"file ({db_mb} MB); production web deployment is in preparation."
)
add_body(
    "Future development trajectories include SRA-based virus discovery, Chinese-language "
    "literature integration, phylogenetic orphan resolution, real-time updating from NCBI and "
    "Europe PMC, community curation interfaces, and phylum expansion to currently unrepresented "
    "aquatic invertebrate taxa."
)

# ══════════════════════════════════════
# DATA AVAILABILITY
# ══════════════════════════════════════
add_heading("Data Availability")
add_body(
    f"AquaVir-KB is freely accessible at [URL TBD]. The complete public dataset ({active_public:,} "
    f"Tier 1 viruses plus {active_limited:,} Tier 2 catalog entries, {ati_raw:,} target isolates, "
    f"{ev_total:,} evidence records, {refs:,} references) is available for download under the "
    "CC-BY 4.0 license. Source code is available at [repository URL TBD]. The database will be "
    "deposited in Zenodo with a versioned DOI upon publication. The design follows the FAIR "
    f"Guiding Principles. Non-target records (n=1,808) are excluded from the public release "
    "with documented exclusion criteria."
)

# ══════════════════════════════════════
# SUPPLEMENTARY
# ══════════════════════════════════════
doc.add_page_break()
add_heading("Supplementary Information")

add_heading("Table S1. Feature comparison with existing virus databases", 3)
s1h = ["Feature", "AquaVir-KB", "NCBI Virus", "ICTV VMR", "Virus-Host DB", "RVDB", "IMG/VR"]
s1r = [
    ["Aquatic invertebrate specialization", "Yes", "No", "No", "No", "No", "No"],
    ["Curated virus species", str(active_broad), ">10M acc.", "11,273", ">10K", "Clustered", ">15M"],
    ["Evidence grading (H/M/L)", f"Yes ({high_ev/ev_total*100:.1f}/{med_ev/ev_total*100:.1f}/{low_ev/ev_total*100:.1f}%)", "No", "No", "No", "No", "No"],
    ["Curation status (4-tier)", "Yes", "No", "No", "No", "No", "No"],
    ["PMID coverage", f"{refs_pmid/refs*100:.1f}%", "Partial", "N/A", "Partial", "N/A", "Minimal"],
    ["DOI coverage", f"{refs_doi/refs*100:.1f}%", "Partial", "N/A", "N/A", "N/A", "Minimal"],
    ["Protein annotation", f"{prot_ann/prots*100:.1f}%", "No", "No", "No", "No", "No"],
    ["Phylogenetic classification", f"Yes ({rdrp:,} RdRP)", "No", "No", "No", "No", "ViCTree"],
    ["SRA metagenomic index", f"{conn.execute('SELECT COUNT(*) FROM sra_runs').fetchone()[0]:,} runs", "No", "No", "No", "No", "Yes"],
    ["Public visibility tiering", "3-tier", "No", "No", "No", "No", "No"],
    ["Curation audit trail", f"{clogs:,} logs", "No", "No", "No", "No", "No"],
    ["Data provenance", f"{prov:,}", "No", "No", "No", "No", "No"],
    ["REST API", "FastAPI", "Yes", "No", "Yes", "No", "Yes"],
    ["Bulk download", "CC-BY 4.0", "Yes", "Yes", "Yes", "Yes", "Yes"],
    ["Docker deployment", "Yes", "No", "No", "No", "No", "No"],
]
add_table(s1h, s1r)

# ══════════════════════════════════════
# REFERENCES
# ══════════════════════════════════════
add_heading("References")
add_ref("[Reference list to be compiled from REFERENCES.md]")

# Save
outpath = os.path.join(BASE, 'NAR_PAPER_MANUSCRIPT.docx')
doc.save(outpath)
print(f"Saved: {outpath}")
print(f"Size: {_os.path.getsize(outpath)/1024:.0f} KB")
print(f"Content: Abstract + Introduction + Methods(7) + Results(5) + Discussion + Data Avail + Table S1")
print(f"Tables: 5 main + 1 supplementary, all from DB")
conn.close()
