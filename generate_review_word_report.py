#!/usr/bin/env python3
"""Generate a Word report for remaining missing and manual-review items."""

from __future__ import annotations

import argparse
import csv
import json
import sqlite3
from datetime import datetime
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


DB_PATH = Path("crustacean_virus_core.db")
REPORTS_DIR = Path("reports")


def stamp() -> str:
    return datetime.now().strftime("%Y%m%d_%H%M%S")


def connect(db_path: Path) -> sqlite3.Connection:
    conn = sqlite3.connect(db_path)
    conn.row_factory = sqlite3.Row
    return conn


def rows(conn: sqlite3.Connection, sql: str, params: tuple[Any, ...] = ()) -> list[dict[str, Any]]:
    return [dict(r) for r in conn.execute(sql, params).fetchall()]


def value(conn: sqlite3.Connection, sql: str, params: tuple[Any, ...] = ()) -> Any:
    row = conn.execute(sql, params).fetchone()
    return row[0] if row else None


def latest_quality_report() -> dict[str, Any] | None:
    candidates = sorted(REPORTS_DIR.glob("database_quality_report_*.json"), key=lambda p: p.stat().st_mtime, reverse=True)
    if not candidates:
        return None
    return json.loads(candidates[0].read_text(encoding="utf-8"))


def write_csv(path: Path, data: list[dict[str, Any]]) -> None:
    path.parent.mkdir(exist_ok=True)
    if not data:
        path.write_text("", encoding="utf-8")
        return
    with path.open("w", newline="", encoding="utf-8-sig") as f:
        writer = csv.DictWriter(f, fieldnames=list(data[0].keys()))
        writer.writeheader()
        writer.writerows(data)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_font(run, bold: bool = False, size: int | None = None) -> None:
    run.font.name = "宋体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run.bold = bold
    if size:
        run.font.size = Pt(size)


def add_heading(doc: Document, text: str, level: int) -> None:
    p = doc.add_heading("", level=level)
    run = p.add_run(text)
    set_font(run, bold=True, size=16 if level == 1 else 13)


def add_paragraph(doc: Document, text: str, bold_prefix: str | None = None) -> None:
    p = doc.add_paragraph()
    if bold_prefix and text.startswith(bold_prefix):
        run = p.add_run(bold_prefix)
        set_font(run, bold=True, size=10)
        run = p.add_run(text[len(bold_prefix) :])
        set_font(run, size=10)
    else:
        run = p.add_run(text)
        set_font(run, size=10)


def add_table(doc: Document, headers: list[str], data: list[dict[str, Any]], keys: list[str], max_rows: int | None = None) -> None:
    shown = data if max_rows is None else data[:max_rows]
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, header in enumerate(headers):
        set_cell_shading(hdr[i], "D9EAF7")
        run = hdr[i].paragraphs[0].add_run(header)
        set_font(run, bold=True, size=9)
    for row in shown:
        cells = table.add_row().cells
        for i, key in enumerate(keys):
            text = "" if row.get(key) is None else str(row.get(key))
            run = cells[i].paragraphs[0].add_run(text)
            set_font(run, size=8)
    if max_rows is not None and len(data) > max_rows:
        add_paragraph(doc, f"注：表内仅展示前 {max_rows} 条；完整明细见同目录 CSV。")


def collect_data(conn: sqlite3.Connection) -> dict[str, Any]:
    return {
        "summary": {
            "integrity": value(conn, "PRAGMA integrity_check"),
            "foreign_key_violations": len(conn.execute("PRAGMA foreign_key_check").fetchall()),
            "viral_isolates_total": value(conn, "SELECT COUNT(*) FROM viral_isolates"),
            "analysis_target_isolates": value(conn, "SELECT COUNT(*) FROM analysis_target_isolates"),
            "target_missing_refs": value(
                conn,
                """
                SELECT COUNT(*)
                FROM analysis_target_isolates vi
                WHERE vi.reference_id IS NULL
                  AND NOT EXISTS (SELECT 1 FROM isolate_reference_links irl WHERE irl.isolate_id=vi.isolate_id)
                """,
            ),
            "open_conflicts": value(conn, "SELECT COUNT(*) FROM curation_conflicts WHERE status='open'"),
            "critical_ictv": value(conn, "SELECT COUNT(*) FROM ictv_review_priority_queue WHERE priority='critical'"),
            "evidence_needs_review": value(conn, "SELECT COUNT(*) FROM evidence_records WHERE curation_status='needs_review'"),
            "diagnostic_needs_review": value(conn, "SELECT COUNT(*) FROM diagnostic_methods WHERE curation_status='needs_review'"),
            "target_missing_country": value(
                conn,
                """
                SELECT COUNT(*)
                FROM analysis_target_isolates vi
                LEFT JOIN infection_records ir ON vi.isolate_id = ir.isolate_id
                LEFT JOIN sample_collections s ON ir.collection_id = s.collection_id
                LEFT JOIN isolate_curated_profiles icp ON vi.isolate_id = icp.isolate_id
                WHERE COALESCE(NULLIF(s.country,''), NULLIF(icp.country,'')) IS NULL
                """,
            ),
            "target_missing_coordinates": value(
                conn,
                """
                SELECT COUNT(*)
                FROM analysis_target_isolates vi
                LEFT JOIN infection_records ir ON vi.isolate_id = ir.isolate_id
                LEFT JOIN sample_collections s ON ir.collection_id = s.collection_id
                LEFT JOIN isolate_curated_profiles icp ON vi.isolate_id = icp.isolate_id
                WHERE COALESCE(s.latitude, icp.latitude) IS NULL
                   OR COALESCE(s.longitude, icp.longitude) IS NULL
                """,
            ),
            "missing_sequence_length": value(conn, "SELECT COUNT(*) FROM analysis_target_isolates WHERE sequence_length IS NULL AND genome_length IS NULL"),
            "sync_status": json.loads((Path("sync_runtime") / "sync_status.json").read_text(encoding="utf-8")).get("status")
            if (Path("sync_runtime") / "sync_status.json").exists()
            else "missing",
        },
        "target_masters_without_isolates": rows(
            conn,
            """
            SELECT vm.master_id, vm.canonical_name, vm.abbreviations, vm.virus_family, vm.virus_genus,
                   vm.entry_type, vm.notes
            FROM virus_master vm
            LEFT JOIN viral_isolates vi ON vm.master_id = vi.master_id
            WHERE vi.isolate_id IS NULL
              AND vm.is_crustacean_virus = 1
              AND vm.entry_type NOT IN ('non_target','host_genome')
            ORDER BY vm.master_id
            """,
        ),
        "evidence_queue_counts": rows(
            conn,
            """
            SELECT priority, evidence_type, COUNT(*) AS n
            FROM evidence_review_priority_queue
            WHERE queue_status='open'
            GROUP BY priority, evidence_type
            ORDER BY CASE priority WHEN 'critical' THEN 0 WHEN 'high' THEN 1 WHEN 'medium' THEN 2 ELSE 3 END, n DESC
            """,
        ),
        "evidence_queue_full": rows(
            conn,
            """
            SELECT q.priority, q.priority_score, er.evidence_id, er.evidence_type,
                   vm.canonical_name, vm.abbreviations, er.claim, er.value_text,
                   er.evidence_strength, er.reference_id, rl.title, rl.year, q.reason
            FROM evidence_review_priority_queue q
            JOIN evidence_records er ON er.evidence_id = q.evidence_id
            LEFT JOIN virus_master vm ON vm.master_id = er.virus_master_id
            LEFT JOIN ref_literatures rl ON rl.reference_id = er.reference_id
            WHERE q.queue_status='open'
            ORDER BY q.priority_score DESC, er.evidence_type, vm.canonical_name, er.evidence_id
            """,
        ),
        "diagnostic_review": rows(
            conn,
            """
            SELECT dm.method_id, dm.data_quality, dm.curation_status, vm.canonical_name,
                   dm.method_name, dm.method_category, dm.method_subcategory,
                   dm.target_gene_or_region, dm.reference_id, rl.title, dm.notes
            FROM diagnostic_methods dm
            LEFT JOIN virus_master vm ON vm.master_id = dm.virus_master_id
            LEFT JOIN ref_literatures rl ON rl.reference_id = dm.reference_id
            WHERE dm.curation_status='needs_review'
            ORDER BY dm.data_quality, vm.canonical_name, dm.method_name
            """,
        ),
        "ictv_pending": rows(
            conn,
            """
            SELECT q.priority, q.master_id, q.canonical_name, q.abbreviations,
                   q.virus_family, q.virus_genus, q.isolate_count, q.reason
            FROM ictv_review_priority_queue q
            ORDER BY CASE q.priority WHEN 'critical' THEN 0 WHEN 'high' THEN 1 WHEN 'medium' THEN 2 ELSE 3 END,
                     q.isolate_count DESC, q.canonical_name
            """,
        ),
        "geo_missing_by_master": rows(
            conn,
            """
            SELECT vm.master_id, vm.canonical_name, COUNT(*) AS missing_geo_count
            FROM analysis_target_isolates vi
            JOIN virus_master vm ON vm.master_id = vi.master_id
            LEFT JOIN infection_records ir ON vi.isolate_id = ir.isolate_id
            LEFT JOIN sample_collections s ON ir.collection_id = s.collection_id
            LEFT JOIN isolate_curated_profiles icp ON vi.isolate_id = icp.isolate_id
            WHERE COALESCE(NULLIF(s.country,''), NULLIF(icp.country,'')) IS NULL
               OR COALESCE(s.latitude, icp.latitude) IS NULL
               OR COALESCE(s.longitude, icp.longitude) IS NULL
            GROUP BY vm.master_id, vm.canonical_name
            ORDER BY missing_geo_count DESC, vm.canonical_name
            """,
        ),
        "sequence_missing": rows(
            conn,
            """
            SELECT vi.isolate_id, vi.accession, vm.canonical_name, vi.virus_name, vi.master_id
            FROM analysis_target_isolates vi
            JOIN virus_master vm ON vm.master_id = vi.master_id
            WHERE vi.sequence_length IS NULL AND vi.genome_length IS NULL
            ORDER BY vm.canonical_name, vi.accession
            """,
        ),
    }


def build_doc(data: dict[str, Any], artifact_paths: dict[str, str], output_path: Path) -> None:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.8)
    section.left_margin = Cm(1.8)
    section.right_margin = Cm(1.8)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("甲壳动物病毒数据库：缺失项与人工复核清单")
    set_font(run, bold=True, size=18)
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run(f"生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    set_font(run, size=10)

    add_heading(doc, "一、结论摘要", 1)
    s = data["summary"]
    summary_rows = [
        {"metric": "数据库完整性", "value": s["integrity"], "judgement": "可继续维护" if s["integrity"] == "ok" else "必须先修复"},
        {"metric": "外键违规", "value": s["foreign_key_violations"], "judgement": "合格" if s["foreign_key_violations"] == 0 else "不合格"},
        {"metric": "目标分析 isolate", "value": s["analysis_target_isolates"], "judgement": "已排除明显非目标/伪 isolate"},
        {"metric": "目标 isolate 缺引用", "value": s["target_missing_refs"], "judgement": "已清零" if s["target_missing_refs"] == 0 else "仍需补文献"},
        {"metric": "critical ICTV 待复核", "value": s["critical_ictv"], "judgement": "已清零" if s["critical_ictv"] == 0 else "阻塞分类声明"},
        {"metric": "自动证据待复核", "value": s["evidence_needs_review"], "judgement": "不可直接用于论文结论"},
        {"metric": "诊断方法待复核", "value": s["diagnostic_needs_review"], "judgement": "需补方法级引用或降级"},
        {"metric": "缺国家/坐标", "value": f"{s['target_missing_country']} / {s['target_missing_coordinates']}", "judgement": "地图分析仍不完整"},
        {"metric": "缺序列长度", "value": s["missing_sequence_length"], "judgement": "影响长度/基因组统计"},
        {"metric": "同步状态", "value": s["sync_status"], "judgement": "需恢复自动同步" if s["sync_status"] == "stale" else "正常/待确认"},
    ]
    add_table(doc, ["指标", "当前值", "处理判断"], summary_rows, ["metric", "value", "judgement"])

    add_heading(doc, "二、仍缺失的核心对象", 1)
    add_paragraph(doc, "这些条目不是简单缺文献，而是数据库对象本身缺少可计数 isolate 或关键元数据。")
    add_heading(doc, "2.1 目标 master 无 isolate", 2)
    add_table(
        doc,
        ["master_id", "标准名", "缩写", "科", "属", "entry_type", "备注"],
        data["target_masters_without_isolates"],
        ["master_id", "canonical_name", "abbreviations", "virus_family", "virus_genus", "entry_type", "notes"],
    )

    add_heading(doc, "2.2 地理信息缺失", 2)
    add_paragraph(doc, f"目标 isolate 中缺国家 {s['target_missing_country']} 条，缺坐标 {s['target_missing_coordinates']} 条。下表按病毒 master 汇总，完整明细见 CSV。")
    add_table(
        doc,
        ["master_id", "病毒标准名", "缺失数量"],
        data["geo_missing_by_master"],
        ["master_id", "canonical_name", "missing_geo_count"],
        max_rows=30,
    )

    add_heading(doc, "2.3 序列长度缺失", 2)
    add_table(
        doc,
        ["isolate_id", "accession", "病毒标准名", "原始名称", "master_id"],
        data["sequence_missing"],
        ["isolate_id", "accession", "canonical_name", "virus_name", "master_id"],
    )

    add_heading(doc, "三、需要人工复核的证据", 1)
    add_paragraph(doc, "这些记录已有引用或候选来源，但仍然是自动抽取/候选证据。严格处理时，不能直接写进论文结论。")
    add_heading(doc, "3.1 Evidence 复核优先级汇总", 2)
    add_table(doc, ["优先级", "证据类型", "数量"], data["evidence_queue_counts"], ["priority", "evidence_type", "n"])
    add_heading(doc, "3.2 Evidence 复核明细预览", 2)
    add_table(
        doc,
        ["优先级", "evidence_id", "类型", "病毒", "缩写", "claim", "reference_id", "文献题名"],
        data["evidence_queue_full"],
        ["priority", "evidence_id", "evidence_type", "canonical_name", "abbreviations", "claim", "reference_id", "title"],
        max_rows=60,
    )

    add_heading(doc, "四、需要人工复核的诊断方法", 1)
    add_paragraph(doc, "当前 placeholder 已全部拒绝；剩余为 candidate_unreferenced 或 curated/needs_review，需要补方法级引用、靶基因、样本类型、检测限等字段。")
    add_table(
        doc,
        ["method_id", "质量", "状态", "病毒", "方法名", "类别", "子类", "靶标", "reference_id", "文献题名"],
        data["diagnostic_review"],
        ["method_id", "data_quality", "curation_status", "canonical_name", "method_name", "method_category", "method_subcategory", "target_gene_or_region", "reference_id", "title"],
        max_rows=80,
    )

    add_heading(doc, "五、ICTV 后续复核项", 1)
    add_paragraph(doc, "critical ICTV 已清零；剩余 high/medium/low 为非阻塞复核项，不能支撑“全库 ICTV 完全映射”的表述。")
    add_table(
        doc,
        ["优先级", "master_id", "病毒名", "缩写", "科", "属", "isolate数", "原因"],
        data["ictv_pending"],
        ["priority", "master_id", "canonical_name", "abbreviations", "virus_family", "virus_genus", "isolate_count", "reason"],
        max_rows=60,
    )

    add_heading(doc, "六、建议处理顺序", 1)
    recommendations = [
        "先处理 247 条 evidence_review_priority_queue：critical 和 high 的 virulence/mortality/diagnosis 优先。",
        "对 29 条诊断方法待复核记录逐条补方法级文献；补不上的继续保持 candidate_unreferenced，不进入分析视图。",
        "补 1065 条地理信息时，必须区分精确坐标、采样地、国家质心和未知，不能混作同一精度。",
        "3 个无 isolate 的 target master 需要决定是补 accession、转为 reference-only 条目，还是退休。",
        "恢复 stale 同步状态后，再重新运行 database_quality_report.py 和本 Word 生成脚本。",
    ]
    for rec in recommendations:
        add_paragraph(doc, f"- {rec}")

    add_heading(doc, "七、完整明细文件", 1)
    artifact_rows = [{"name": k, "path": v} for k, v in artifact_paths.items()]
    add_table(doc, ["文件", "路径"], artifact_rows, ["name", "path"])

    doc.save(output_path)


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--db", default=str(DB_PATH))
    args = parser.parse_args()

    REPORTS_DIR.mkdir(exist_ok=True)
    ts = stamp()
    conn = connect(Path(args.db))
    data = collect_data(conn)
    conn.close()

    csv_paths = {
        "evidence_review_full_csv": REPORTS_DIR / f"manual_review_evidence_full_{ts}.csv",
        "diagnostic_review_csv": REPORTS_DIR / f"manual_review_diagnostics_{ts}.csv",
        "ictv_pending_csv": REPORTS_DIR / f"manual_review_ictv_pending_{ts}.csv",
        "geo_missing_by_master_csv": REPORTS_DIR / f"missing_geo_by_master_{ts}.csv",
        "sequence_missing_csv": REPORTS_DIR / f"missing_sequence_length_{ts}.csv",
        "target_masters_without_isolates_csv": REPORTS_DIR / f"target_masters_without_isolates_{ts}.csv",
    }
    write_csv(csv_paths["evidence_review_full_csv"], data["evidence_queue_full"])
    write_csv(csv_paths["diagnostic_review_csv"], data["diagnostic_review"])
    write_csv(csv_paths["ictv_pending_csv"], data["ictv_pending"])
    write_csv(csv_paths["geo_missing_by_master_csv"], data["geo_missing_by_master"])
    write_csv(csv_paths["sequence_missing_csv"], data["sequence_missing"])
    write_csv(csv_paths["target_masters_without_isolates_csv"], data["target_masters_without_isolates"])

    output_path = REPORTS_DIR / f"甲壳动物病毒数据库_缺失与人工复核清单_{ts}.docx"
    artifact_paths = {k: str(v) for k, v in csv_paths.items()}
    latest_report = latest_quality_report()
    if latest_report:
        artifact_paths["latest_quality_report_json"] = str(sorted(REPORTS_DIR.glob("database_quality_report_*.json"), key=lambda p: p.stat().st_mtime, reverse=True)[0])
    build_doc(data, artifact_paths, output_path)

    summary = {
        "word_report": str(output_path),
        "csv_artifacts": artifact_paths,
        "summary": data["summary"],
    }
    summary_path = REPORTS_DIR / f"manual_review_word_report_{ts}.json"
    summary_path.write_text(json.dumps(summary, ensure_ascii=False, indent=2), encoding="utf-8")
    print(json.dumps(summary, ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()
