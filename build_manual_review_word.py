#!/usr/bin/env python3
"""Build a Word manual-review workbook for unresolved curation items."""

from __future__ import annotations

import sqlite3
from datetime import datetime
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt


DB_PATH = Path("crustacean_virus_core.db")
OUT_DIR = Path("reports")


def connect() -> sqlite3.Connection:
    conn = sqlite3.connect(DB_PATH)
    conn.row_factory = sqlite3.Row
    conn.execute("PRAGMA foreign_keys = ON")
    return conn


def rows(conn: sqlite3.Connection, sql: str, params: tuple[Any, ...] = ()) -> list[dict[str, Any]]:
    return [dict(r) for r in conn.execute(sql, params).fetchall()]


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def add_table(doc: Document, columns: list[str], data: list[dict[str, Any]], title: str | None = None) -> None:
    if title:
        add_heading(doc, title, 2)
    table = doc.add_table(rows=1, cols=len(columns))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, col in enumerate(columns):
        hdr[i].text = col
    for item in data:
        cells = table.add_row().cells
        for i, col in enumerate(columns):
            val = item.get(col, "")
            if val is None:
                val = ""
            text = str(val)
            if len(text) > 500:
                text = text[:497] + "..."
            cells[i].text = text
    doc.add_paragraph()


def set_doc_style(doc: Document) -> None:
    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(9)
    for section in doc.sections:
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width, section.page_height = section.page_height, section.page_width
        section.top_margin = Cm(1.2)
        section.bottom_margin = Cm(1.2)
        section.left_margin = Cm(1.2)
        section.right_margin = Cm(1.2)


def main() -> None:
    OUT_DIR.mkdir(exist_ok=True)
    ts = datetime.now().strftime("%Y%m%d_%H%M%S")
    out_path = OUT_DIR / f"manual_review_workbook_{ts}.docx"

    with connect() as conn:
        summary = [
            {"类别": "证据记录", "数量": conn.execute("SELECT COUNT(*) FROM evidence_records WHERE curation_status='needs_review'").fetchone()[0], "优先级": "P0", "人工判断": "判断原文是否支持致病性/宿主范围/诊断/环境证据"},
            {"类别": "诊断方法", "数量": conn.execute("SELECT COUNT(*) FROM diagnostic_methods WHERE curation_status='needs_review' AND data_quality <> 'placeholder'").fetchone()[0], "优先级": "P0", "人工判断": "确认方法、靶标、验证场景和参考文献"},
            {"类别": "ICTV pending", "数量": conn.execute("SELECT COUNT(*) FROM virus_ictv_status WHERE ictv_status='pending_review'").fetchone()[0], "优先级": "P0", "人工判断": "确认是否正式分类、是否应映射/拒绝/保留未分类"},
            {"类别": "无 isolate 的 target master", "数量": conn.execute("""SELECT COUNT(*) FROM virus_master vm LEFT JOIN viral_isolates vi ON vi.master_id=vm.master_id WHERE vi.isolate_id IS NULL AND vm.is_crustacean_virus=1 AND vm.entry_type NOT IN ('non_target','host_genome')""").fetchone()[0], "优先级": "P0", "人工判断": "决定补 isolate、合并 master、降级或删除"},
            {"类别": "宿主范围复核", "数量": conn.execute("SELECT COUNT(*) FROM auto_host_scope_worklist").fetchone()[0], "优先级": "P1", "人工判断": "确认技术宿主、非甲壳动物、非物种级宿主是否排除"},
            {"类别": "目标株缺宿主", "数量": conn.execute("SELECT COUNT(*) FROM auto_completeness_worklist WHERE issue_type='missing_host'").fetchone()[0], "优先级": "P1", "人工判断": "从 GenBank/BioSample/文献确认宿主"},
            {"类别": "目标株缺地理", "数量": conn.execute("SELECT COUNT(*) FROM auto_completeness_worklist WHERE issue_type='missing_country'").fetchone()[0], "优先级": "P1", "人工判断": "确认国家/坐标来源和精度"},
            {"类别": "目标株缺 genome_type", "数量": conn.execute("SELECT COUNT(*) FROM auto_completeness_worklist WHERE issue_type='missing_genome_type'").fetchone()[0], "优先级": "P2", "人工判断": "按 ICTV/NCBI/文献统一 genome_type"},
        ]

        evidence = rows(
            conn,
            """
            SELECT
                er.evidence_id AS 记录ID,
                COALESCE(vm.canonical_name, '') AS 病毒,
                COALESCE(ch.scientific_name, '') AS 宿主,
                er.evidence_type AS 证据类型,
                er.evidence_strength AS 当前强度,
                er.curation_status AS 当前状态,
                COALESCE(rl.pmid, '') AS PMID,
                COALESCE(rl.doi, '') AS DOI,
                COALESCE(rl.title, '') AS 参考文献标题,
                COALESCE(er.claim, er.value_text, er.context, er.notes, '') AS 待核对原文片段,
                '读原文；确认对象是否为该病毒+该宿主；确认实验/观察是否直接支持证据类型；填写 manual_checked 或 rejected，并记录理由' AS 审核动作
            FROM evidence_records er
            LEFT JOIN virus_master vm ON vm.master_id = er.virus_master_id
            LEFT JOIN crustacean_hosts ch ON ch.host_id = er.host_id
            LEFT JOIN ref_literatures rl ON rl.reference_id = er.reference_id
            WHERE er.curation_status = 'needs_review'
            ORDER BY
                CASE er.evidence_strength WHEN 'high' THEN 0 WHEN 'medium' THEN 1 ELSE 2 END,
                er.evidence_type,
                vm.canonical_name
            LIMIT 80
            """,
        )

        diagnostics = rows(
            conn,
            """
            SELECT
                dm.method_id AS 记录ID,
                COALESCE(vm.canonical_name, '') AS 病毒,
                dm.method_category AS 方法类别,
                dm.method_name AS 方法名称,
                COALESCE(dm.target_gene_or_region, '') AS 靶标,
                COALESCE(dm.sample_type, '') AS 样本类型,
                COALESCE(dm.detection_limit, '') AS 检出限,
                dm.evidence_strength AS 当前强度,
                dm.data_quality AS 数据质量,
                COALESCE(rl.pmid, '') AS PMID,
                COALESCE(rl.doi, '') AS DOI,
                COALESCE(rl.title, '') AS 参考文献标题,
                '确认是否真实检测该病毒；靶标是否明确；是否有验证样本/灵敏度/特异性；合格后 manual_checked，不合格 rejected' AS 审核动作
            FROM diagnostic_methods dm
            LEFT JOIN virus_master vm ON vm.master_id = dm.virus_master_id
            LEFT JOIN ref_literatures rl ON rl.reference_id = dm.reference_id
            WHERE dm.curation_status = 'needs_review'
              AND dm.data_quality <> 'placeholder'
            ORDER BY dm.data_quality, vm.canonical_name
            """,
        )

        ictv = rows(
            conn,
            """
            SELECT
                vm.master_id AS MasterID,
                vm.canonical_name AS 病毒,
                COALESCE(vm.abbreviations, '') AS 缩写,
                COALESCE(vm.virus_family, '') AS 当前科,
                COALESCE(vm.virus_genus, '') AS 当前属,
                vis.ictv_status AS ICTV状态,
                COALESCE(vis.reason, '') AS 原因,
                COALESCE(iq.priority, '') AS 优先级,
                COALESCE(iq.reason, '') AS 队列原因,
                '查 ICTV MSL/VMR、NCBI Taxonomy、原始论文；决定 mapped/rejected/unclassified_not_expected/pending_review，并记录依据' AS 审核动作
            FROM virus_master vm
            JOIN virus_ictv_status vis ON vis.master_id = vm.master_id
            LEFT JOIN ictv_review_priority_queue iq ON iq.master_id = vm.master_id
            WHERE vis.ictv_status = 'pending_review'
            ORDER BY CASE iq.priority WHEN 'critical' THEN 0 WHEN 'high' THEN 1 WHEN 'medium' THEN 2 ELSE 3 END, vm.canonical_name
            """,
        )

        orphan_master = rows(
            conn,
            """
            SELECT
                vm.master_id AS MasterID,
                vm.canonical_name AS 病毒,
                COALESCE(vm.virus_family, '') AS 当前科,
                COALESCE(vm.genome_type, '') AS 基因组类型,
                vm.entry_type AS 条目类型,
                COALESCE(vmq.severity, '') AS 严重性,
                COALESCE(vmq.reason, '') AS 原因,
                '检查是否有遗漏 accession；若是重复则合并；若非目标则改 entry_type；若真实目标则补 isolate' AS 审核动作
            FROM virus_master vm
            LEFT JOIN viral_isolates vi ON vi.master_id = vm.master_id
            LEFT JOIN virus_master_review_queue vmq ON vmq.master_id = vm.master_id
            WHERE vi.isolate_id IS NULL
              AND vm.is_crustacean_virus = 1
              AND vm.entry_type NOT IN ('non_target', 'host_genome')
            ORDER BY vm.canonical_name
            """,
        )

        host_scope = rows(
            conn,
            """
            SELECT
                host_id AS 宿主ID,
                scientific_name AS 宿主名,
                host_type AS 当前宿主类型,
                host_group AS 宿主组,
                COALESCE(taxon_order, '') AS 目,
                issue_type AS 问题类型,
                suggested_scope_status AS 建议范围状态,
                suggested_exclude_from_target_stats AS 是否排除目标统计,
                '确认该名称是自然宿主、技术宿主、非甲壳动物，还是非物种级泛称；必要时查 NCBI Taxonomy/WoRMS/原文' AS 审核动作
            FROM auto_host_scope_worklist
            ORDER BY issue_type, scientific_name
            """,
        )

        missing_host = rows(
            conn,
            """
            SELECT
                accession AS Accession,
                virus_name AS 病毒,
                issue_type AS 问题,
                suggested_source AS 建议来源,
                suggested_action AS 建议动作,
                '查 GenBank source.host、BioSample organism/host、论文 Methods；确认后填 host_id/host_scientific_name，并记录来源' AS 审核动作
            FROM auto_completeness_worklist
            WHERE issue_type = 'missing_host'
            ORDER BY virus_name, accession
            LIMIT 80
            """,
        )

        missing_geo = rows(
            conn,
            """
            SELECT
                accession AS Accession,
                virus_name AS 病毒,
                issue_type AS 问题,
                suggested_source AS 建议来源,
                suggested_action AS 建议动作,
                '先补国家；坐标只能用原文/GenBank明确坐标或标记为国家/省级质心；必须填写 location_precision/coordinates_source' AS 审核动作
            FROM auto_completeness_worklist
            WHERE issue_type IN ('missing_country', 'missing_coordinates')
            ORDER BY virus_name, accession, issue_type
            LIMIT 100
            """,
        )

        genome_type = rows(
            conn,
            """
            SELECT
                accession AS Accession,
                virus_name AS 病毒,
                issue_type AS 问题,
                suggested_source AS 建议来源,
                suggested_action AS 建议动作,
                '按 ICTV/NCBI taxonomy/GenBank molecule_type 统一到 dsDNA/ssDNA/+ssRNA/-ssRNA/dsRNA/retro 等受控值' AS 审核动作
            FROM auto_completeness_worklist
            WHERE issue_type = 'missing_genome_type'
            ORDER BY virus_name, accession
            LIMIT 80
            """,
        )

    doc = Document()
    set_doc_style(doc)
    title = doc.add_heading("甲壳动物病毒数据库人工审核工作表", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(f"生成时间：{datetime.now().isoformat(timespec='seconds')}")
    doc.add_paragraph("用途：把自动程序不能可靠判断、必须人工确认的记录集中列出。请不要直接把 needs_review 批量改为 manual_checked；必须逐条核对来源。")

    add_heading(doc, "一、审核总原则", 1)
    principles = [
        "先 P0 后 P1 后 P2：P0 影响数据库结论可信度，优先审核。",
        "每条通过审核的记录必须能追溯到 PMID、DOI、GenBank、BioSample、ICTV MSL/VMR 或原始论文。",
        "证据审核要区分直接证据和间接证据。实验感染、死亡率、组织病理、特异诊断属于强证据；综述转述、关键词命中、自动抽取只能算弱证据。",
        "地理坐标不能混用。真实采样点、城市/省级质心、国家质心、未知位置必须分别标记。",
        "宿主要区分自然宿主、技术宿主、非甲壳动物、非物种级泛称。E. coli/DH10B 等表达宿主不应进入甲壳动物宿主谱统计。",
        "无法确认的记录保持 needs_review 或 unknown，不要为了完整率强行补值。",
    ]
    for p in principles:
        doc.add_paragraph(p, style="List Bullet")

    add_heading(doc, "二、审核后如何更新数据库", 1)
    steps = [
        "在 Word 表格中先填写判断和依据；建议同时在对应 CSV 中保留人工备注。",
        "证据记录：确认支持则将 curation_status 改为 manual_checked；不支持则 rejected；必要时调整 evidence_strength。",
        "诊断方法：确认方法真实且有文献支撑后 manual_checked；靶标/检出限缺失但方法可信时可保留 curated + needs_review，并补 notes。",
        "ICTV：根据 ICTV/NCBI/原文将状态改为 mapped、rejected、non_target 或 unclassified_not_expected；保留 reason。",
        "宿主范围：技术宿主和非甲壳动物设置 exclude_from_target_stats=1；自然甲壳动物宿主保留 target。",
        "地理字段：补 country、latitude、longitude 时同步填写 location_precision 和 coordinates_source。",
    ]
    for s in steps:
        doc.add_paragraph(s, style="List Number")

    add_table(doc, ["类别", "数量", "优先级", "人工判断"], summary, "三、人工审核任务总表")
    add_table(doc, ["记录ID", "病毒", "宿主", "证据类型", "当前强度", "当前状态", "PMID", "DOI", "参考文献标题", "待核对原文片段", "审核动作"], evidence, "四、P0 证据记录审核（前 80 条）")
    add_table(doc, ["记录ID", "病毒", "方法类别", "方法名称", "靶标", "样本类型", "检出限", "当前强度", "数据质量", "PMID", "DOI", "参考文献标题", "审核动作"], diagnostics, "五、P0 诊断方法审核")
    add_table(doc, ["MasterID", "病毒", "缩写", "当前科", "当前属", "ICTV状态", "原因", "优先级", "队列原因", "审核动作"], ictv, "六、P0 ICTV pending 审核")
    add_table(doc, ["MasterID", "病毒", "当前科", "基因组类型", "条目类型", "严重性", "原因", "审核动作"], orphan_master, "七、P0 无 isolate 的 target master")
    add_table(doc, ["宿主ID", "宿主名", "当前宿主类型", "宿主组", "目", "问题类型", "建议范围状态", "是否排除目标统计", "审核动作"], host_scope, "八、P1 宿主范围审核")
    add_table(doc, ["Accession", "病毒", "问题", "建议来源", "建议动作", "审核动作"], missing_host, "九、P1 缺宿主记录（前 80 条）")
    add_table(doc, ["Accession", "病毒", "问题", "建议来源", "建议动作", "审核动作"], missing_geo, "十、P1 缺地理记录（前 100 条）")
    add_table(doc, ["Accession", "病毒", "问题", "建议来源", "建议动作", "审核动作"], genome_type, "十一、P2 缺 genome_type 记录（前 80 条）")

    doc.add_paragraph("完整清单位置：reports/auto_optimization_20260507_084550/completeness_worklist.csv、host_scope_worklist.csv、annotation_gap_worklist.csv。Word 中只放入优先审核样例和全部 P0 小表，避免文档过大。")
    doc.save(out_path)
    print(out_path)


if __name__ == "__main__":
    main()
