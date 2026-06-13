"""Convert NAR_PAPER_DRAFT.md to NAR_PAPER_MANUSCRIPT.docx with basic formatting."""
import re, os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

BASE = os.path.dirname(__file__)
md_path = os.path.join(BASE, 'NAR_PAPER_DRAFT.md')

with open(md_path, 'r', encoding='utf-8') as f:
    text = f.read()

doc = Document()

# Set default font
style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(12)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.5

# Parse markdown and build docx
lines = text.split('\n')
i = 0
in_table = False
in_code = False

while i < len(lines):
    line = lines[i]

    # Skip empty
    if not line.strip():
        i += 1
        continue

    # Headers
    if line.startswith('# '):
        p = doc.add_paragraph()
        run = p.add_run(line[2:].strip())
        run.bold = True
        run.font.size = Pt(16)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif line.startswith('## '):
        p = doc.add_paragraph()
        run = p.add_run(line[3:].strip())
        run.bold = True
        run.font.size = Pt(14)
    elif line.startswith('### '):
        p = doc.add_paragraph()
        run = p.add_run(line[4:].strip())
        run.bold = True
        run.font.size = Pt(13)
    elif line.startswith('**') and ':**' in line:
        p = doc.add_paragraph()
        run = p.add_run(line.strip('* '))
        run.bold = True
        run.font.size = Pt(12)
    elif line.startswith('---'):
        # Horizontal rule - skip
        pass

    # Tables
    elif line.startswith('|') and '|' in line[1:]:
        if not in_table:
            in_table = True
            table_data = []
        # Parse table row
        cells = [c.strip() for c in line.split('|')[1:-1]]
        table_data.append(cells)
        # Check if next line is separator or end
        if i + 1 >= len(lines) or not lines[i+1].startswith('|'):
            # End of table
            if len(table_data) > 1:
                # Remove separator row if present
                if all(re.match(r'^[-:]+$', c) for c in table_data[1] if c):
                    table_data.pop(1)
                num_cols = max(len(row) for row in table_data)
                table = doc.add_table(rows=len(table_data), cols=num_cols)
                table.style = 'Table Grid'
                for ri, row in enumerate(table_data):
                    for ci, cell_text in enumerate(row):
                        if ci < num_cols:
                            cell = table.rows[ri].cells[ci]
                            cell.text = cell_text
                            # Bold header
                            if ri == 0:
                                for paragraph in cell.paragraphs:
                                    for run in paragraph.runs:
                                        run.bold = True
                                        run.font.size = Pt(9)
                            else:
                                for paragraph in cell.paragraphs:
                                    for run in paragraph.runs:
                                        run.font.size = Pt(9)
                doc.add_paragraph()  # space after table
            table_data = []
            in_table = False

    # Code blocks
    elif line.startswith('```'):
        in_code = not in_code
    elif in_code:
        pass  # skip code blocks

    # Regular paragraph
    else:
        # Process inline formatting
        text_line = line.strip()
        if text_line:
            p = doc.add_paragraph()
            # Bold text between ** **
            parts = re.split(r'(\*\*.*?\*\*)', text_line)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = p.add_run(part[2:-2])
                    run.bold = True
                else:
                    # Italic between * *
                    sub_parts = re.split(r'(\*[^*]+\*)', part)
                    for sp in sub_parts:
                        if sp.startswith('*') and sp.endswith('*') and len(sp) > 2:
                            run = p.add_run(sp[1:-1])
                            run.italic = True
                        else:
                            p.add_run(sp)

    i += 1

# Save
docx_path = os.path.join(BASE, 'NAR_PAPER_MANUSCRIPT.docx')
doc.save(docx_path)
print(f"Saved: {docx_path}")
print(f"File size: {os.path.getsize(docx_path)/1024:.0f} KB")
