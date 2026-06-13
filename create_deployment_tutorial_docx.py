from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


OUT_DIR = Path("docs")
OUT_DIR.mkdir(exist_ok=True)
OUT_FILE = OUT_DIR / "CrustaVirus_DB_public_deployment_GitHub_domain_ICP_DOI_tutorial.docx"


def setup_doc() -> Document:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)
    for style_name in ["Normal", "Title", "Heading 1", "Heading 2", "Heading 3"]:
        style = doc.styles[style_name]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    doc.styles["Normal"].font.size = Pt(10.5)
    return doc


doc = setup_doc()


def title(text: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(text)
    run.bold = True
    run.font.size = Pt(20)


def p(text: str = "") -> None:
    doc.add_paragraph(text)


def h1(text: str) -> None:
    doc.add_heading(text, level=1)


def h2(text: str) -> None:
    doc.add_heading(text, level=2)


def bullet(text: str) -> None:
    doc.add_paragraph(text, style="List Bullet")


def num(text: str) -> None:
    doc.add_paragraph(text, style="List Number")


def code(text: str) -> None:
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(text)
    run.font.name = "Consolas"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
    run.font.size = Pt(9)
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), "F2F2F2")
    paragraph._p.get_or_add_pPr().append(shading)


def table(headers: list[str], rows: list[list[str]]) -> None:
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    for idx, value in enumerate(headers):
        t.rows[0].cells[idx].text = value
    for row in rows:
        cells = t.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = value
    p()


title("CrustaVirus DB 公网部署、GitHub、域名备案与 DOI 傻瓜教程")
p("版本：v0.1 预发布操作手册")
p("用途：指导你从零开始准备 GitHub 仓库、域名、腾讯云服务器、ICP备案、HTTPS 部署、DOI 归档和 NAR 投稿前检查。")
p("核心原则：基础设施现在就可以做；正式 DOI 和 NAR 投稿必须等数据库冻结、公网可用、维护主体真实后再做。")

h1("0. 先看结论")
h2("现在可以做")
for item in [
    "建 GitHub 仓库，先放代码、脚本、模板和文档骨架。",
    "买域名，最好用 PI、课题组或机构能长期控制的账号。",
    "买腾讯云中国大陆服务器并启动 ICP 备案。备案和数据库完善可以并行。",
    "部署 staging 站点，例如 https://staging.crustavirusdb.org，并标注 Pre-release / under curation。",
    "继续跑 release_gate.py、nar_readiness_check.py，并人工核对证据队列。",
]:
    bullet(item)

h2("现在不要做")
for item in [
    "不要发布正式 v1.0.0 DOI。",
    "不要写 NAR-ready、fully curated、final release。",
    "不要把 crustacean_virus_core.db 直接上传到普通 GitHub 仓库。",
    "不要上传 notification_config.json、.env、API key、邮箱授权码、服务器密码。",
    "不要用 ngrok、frp、个人电脑端口映射作为 NAR 公网 URL。",
]:
    bullet(item)

h1("1. 总体路线")
code(
    "GitHub 仓库\n"
    "  - 放代码、模板、部署脚本、README、LICENSE、CITATION.cff\n"
    "  - 不放数据库大文件，不放任何密钥\n\n"
    "腾讯云服务器 + 域名\n"
    "  - 跑 FastAPI 后端、SQLite 数据库、public_downloads\n"
    "  - Nginx 反向代理 + HTTPS\n"
    "  - 对 NAR 提供真实公网 URL\n\n"
    "Zenodo / Figshare\n"
    "  - 放 v1.0.0 release bundle\n"
    "  - 获得 DOI\n"
    "  - 长期归档"
)
table(
    ["事项", "现在能不能做", "预计时间", "投稿前是否必须完成"],
    [
        ["GitHub 仓库", "能", "0.5-1 天", "是"],
        ["买域名", "能", "0.5 天", "是"],
        ["买腾讯云服务器", "能", "0.5 天", "大陆部署必须"],
        ["ICP备案", "建议立刻做", "通常 2-6 周", "大陆服务器必须"],
        ["部署 staging", "能", "1-3 天", "建议"],
        ["数据库人工核对", "继续做", "数天到数周", "取决于论文主张"],
        ["正式 DOI", "先别急", "0.5-2 天", "投稿前必须"],
        ["NAR pre-submission", "现在不能", "全部就绪后", "是"],
    ],
)

h1("2. GitHub 仓库操作")
h2("2.1 注册和建仓库")
for step in [
    "打开 https://github.com/ 并注册账号。",
    "建议开启 2FA 双因素认证。",
    "登录后右上角点击 +，选择 New repository。",
    "Repository name 填 crustavirus-db 或 crustavirus-database。",
    "Description 填 A release-filtered database and web service for crustacean-associated viruses。",
    "数据库未完善时建议先选 Private；准备公开时再改 Public。",
    "勾选 Add a README file。",
    "点击 Create repository。",
]:
    num(step)

h2("2.2 什么能上传，什么不能上传")
table(
    ["文件类型", "是否上传 GitHub", "说明"],
    [
        ["backend.py / api_models.py / scripts", "可以", "上传前检查是否含密钥"],
        ["templates/", "可以", "网页模板可以上传"],
        ["requirements.txt / environment.yml", "可以", "复现环境需要"],
        ["README.md / LICENSE / CITATION.cff", "可以", "正式公开前必须清理 placeholder"],
        ["crustacean_virus_core.db", "不要放普通仓库", "约 235 MB，不适合普通 GitHub 仓库"],
        ["public_downloads 大文件", "不建议", "正式版放 Zenodo/Figshare 或 GitHub Release"],
        ["notification_config.json", "绝对不要", "可能含邮箱授权码"],
        [".env / API key / 服务器密码", "绝对不要", "泄露后必须吊销"],
        ["backups/ / *.db-wal / *.db-shm", "不要", "备份和 SQLite 临时文件不进仓库"],
    ],
)

h2("2.3 .gitignore 至少包含")
code(
    "*.db\n*.db-wal\n*.db-shm\nbackups/\nmaintenance_archive/internal_candidate_exports/\n"
    "notification_config.json\n.env\n*.log\n__pycache__/\n.venv/\n"
    "sequences/\npublic_downloads/*.fasta\npublic_downloads/*.xlsx\ndownloads/exports/*.tsv"
)

h2("2.4 新手推荐：GitHub Desktop")
for step in [
    "下载 GitHub Desktop：https://desktop.github.com/ 。",
    "登录 GitHub 账号。",
    "File -> Add local repository，选择 F:\\甲壳动物数据库。",
    "如果提示不是 Git 仓库，选择 Create a repository。",
    "提交前认真检查 Changes 列表，确认没有 .db、密钥、备份、日志。",
    "Summary 写 Initial code and documentation skeleton。",
    "点击 Commit。",
    "点击 Publish repository 或 Push origin。",
]:
    num(step)

h1("3. 域名注册")
h2("3.1 域名怎么选")
for item in [
    "最好：机构二级域名，例如 crustavirus.xxx.edu.cn。",
    "很好：课题组/PI 长期控制的独立域名，例如 crustavirusdb.org。",
    "不推荐但可用：学生个人名义域名。它对 5 年维护主体的说服力弱。",
]:
    bullet(item)
p("建议候选：crustavirusdb.org、crustavirus.org、crustavirdb.org、crustavirus.cn。")

h2("3.2 腾讯云买域名")
for step in [
    "打开腾讯云控制台：https://console.cloud.tencent.com/ 。",
    "完成实名认证。",
    "进入“域名注册”。",
    "搜索域名，例如 crustavirusdb.org。",
    "选择购买年限，建议至少 3-5 年。",
    "填写域名持有人信息。若最终用单位备案，域名最好由单位/PI/课题组可控账号持有。",
    "付款后进入“域名管理”，完成域名实名认证。",
]:
    num(step)

h2("3.3 DNS 解析")
p("买好服务器并拿到公网 IP 后，在腾讯云 DNS 解析里添加：")
table(
    ["记录类型", "主机记录", "记录值", "用途"],
    [
        ["A", "@", "服务器公网 IP", "crustavirusdb.org"],
        ["A", "www", "服务器公网 IP", "www.crustavirusdb.org"],
        ["A", "staging", "服务器公网 IP", "staging.crustavirusdb.org"],
    ],
)

h1("4. 腾讯云服务器")
h2("4.1 买什么配置")
p("推荐 Ubuntu 22.04 LTS。最低 2 核 4G、80-100 GB 硬盘；更稳妥是 4 核 8G、200 GB 硬盘。")
p("如果要备案，腾讯云官方要求中国境内有公网 IP 的 CVM 或 Lighthouse，包年包月，购买 3 个月及以上，备案期间剩余有效期满足要求。不要买按量付费临时服务器备案。")

h2("4.2 购买步骤")
for step in [
    "进入腾讯云控制台，选择云服务器 CVM 或轻量应用服务器。",
    "地域选择中国大陆，例如广州、上海、北京。",
    "镜像选择 Ubuntu Server 22.04 LTS。",
    "公网 IP 必须开启。",
    "计费选择包年包月，至少 3 个月，建议 1 年。",
    "安全组开放 22、80、443。不要直接把 8000 暴露到公网。",
    "登录方式优先 SSH 密钥；新手可先用密码，但要保存好。",
    "购买后记录公网 IP。",
]:
    num(step)

h1("5. ICP 备案")
h2("5.1 是否必须备案")
table(
    ["部署位置", "是否需要 ICP 备案", "说明"],
    [
        ["腾讯云中国大陆服务器", "需要", "网站开通前必须备案"],
        ["腾讯云香港/海外服务器", "通常不需要", "不能用于中国大陆备案"],
        ["学校已有备案域名", "看机构流程", "可能走学校信息中心"],
    ],
)

h2("5.2 备案材料")
for item in [
    "实名认证的腾讯云账号。",
    "已实名域名。",
    "符合备案条件的中国大陆云服务器。",
    "备案主体信息：个人或单位。NAR 项目建议单位/PI/课题组主体。",
    "负责人身份证件、手机号、邮箱。",
    "网站名称：建议“甲壳动物病毒数据库”。",
    "服务内容：建议“科研数据展示、检索与下载”。",
]:
    bullet(item)

h2("5.3 腾讯云备案流程")
for step in [
    "腾讯云控制台搜索“备案”或“ICP 备案”。",
    "点击开始备案。",
    "填写备案域名，系统识别备案类型。",
    "填写主体信息和网站信息。",
    "上传证件照片、核验材料。",
    "按提示做人脸核验。",
    "提交腾讯云初审。腾讯云官方说明通常 1-2 个工作日。",
    "腾讯云初审通过后，提交管局。",
    "收到工信部短信后，24 小时内完成短信核验。",
    "等待管局审核。腾讯云官方说明管局审核不超过 20 个工作日。",
    "审核通过后，腾讯云会短信/邮件通知。",
]:
    num(step)
p("备案通过并开通网站后，通常还需要在网站开通日起 30 日内提交公安联网备案。")

h1("6. 部署 FastAPI 网站")
h2("6.1 连接服务器")
code("ssh root@你的服务器公网IP")

h2("6.2 安装软件")
code(
    "apt update\n"
    "apt install -y git python3 python3-venv python3-pip nginx ufw certbot python3-certbot-nginx\n"
    "ufw allow OpenSSH\nufw allow 80\nufw allow 443\nufw enable"
)

h2("6.3 创建运行用户和目录")
code("adduser crustavirus\nmkdir -p /opt/crustavirus\nchown -R crustavirus:crustavirus /opt/crustavirus")

h2("6.4 拉取 GitHub 代码")
code(
    "su - crustavirus\n"
    "cd /opt/crustavirus\n"
    "git clone https://github.com/你的账号/crustavirus-db.git app\n"
    "cd app\n"
    "python3 -m venv .venv\n"
    "source .venv/bin/activate\n"
    "pip install -r requirements.txt"
)

h2("6.5 上传数据库和下载文件")
p("数据库不要通过 GitHub 普通仓库传。用 WinSCP、SFTP 或 scp 上传。")
code(
    "scp crustacean_virus_core.db crustavirus@服务器IP:/opt/crustavirus/app/\n"
    "scp -r public_downloads crustavirus@服务器IP:/opt/crustavirus/app/"
)

h2("6.6 配置 systemd 服务")
code("sudo nano /etc/systemd/system/crustavirus.service")
code(
    "[Unit]\nDescription=CrustaVirus DB FastAPI service\nAfter=network.target\n\n"
    "[Service]\nUser=crustavirus\nGroup=crustavirus\nWorkingDirectory=/opt/crustavirus/app\n"
    "Environment=CRUSTAVIRUS_API_KEY=换成很长的随机字符串\n"
    "ExecStart=/opt/crustavirus/app/.venv/bin/uvicorn backend:app --host 127.0.0.1 --port 8000\n"
    "Restart=always\nRestartSec=5\n\n[Install]\nWantedBy=multi-user.target"
)
code("sudo systemctl daemon-reload\nsudo systemctl enable crustavirus\nsudo systemctl start crustavirus\nsudo systemctl status crustavirus")

h2("6.7 配置 Nginx")
code("sudo nano /etc/nginx/sites-available/crustavirus")
code(
    "server {\n"
    "    listen 80;\n"
    "    server_name crustavirusdb.org www.crustavirusdb.org;\n"
    "    client_max_body_size 200M;\n"
    "    location / {\n"
    "        proxy_pass http://127.0.0.1:8000;\n"
    "        proxy_set_header Host $host;\n"
    "        proxy_set_header X-Real-IP $remote_addr;\n"
    "        proxy_set_header X-Forwarded-For $proxy_add_x_forwarded_for;\n"
    "        proxy_set_header X-Forwarded-Proto $scheme;\n"
    "    }\n"
    "}"
)
code("sudo ln -s /etc/nginx/sites-available/crustavirus /etc/nginx/sites-enabled/crustavirus\nsudo nginx -t\nsudo systemctl reload nginx")

h2("6.8 配 HTTPS")
code("sudo certbot --nginx -d crustavirusdb.org -d www.crustavirusdb.org")
p("按提示输入邮箱、同意条款，并选择自动跳转 HTTPS。")

h2("6.9 公网 smoke test")
code(
    "curl -I https://crustavirusdb.org/\n"
    "curl -I https://crustavirusdb.org/api/health\n"
    "curl -I https://crustavirusdb.org/api/stats\n"
    "curl -I https://crustavirusdb.org/download\n"
    "curl -I https://crustavirusdb.org/api/download/all_sequences.fasta"
)
p("这些接口应返回 200 或合理的 HTTPS 跳转，不能跳登录，不能 500。")

h1("7. DOI")
h2("7.1 什么时候做 DOI")
p("不要现在发布正式 DOI。等数据库冻结、release_gate.py 通过、下载包稳定、真实作者机构确定、公网 URL 稳定后再发布。")

h2("7.2 Zenodo 手动上传流程")
for step in [
    "注册并登录 https://zenodo.org/ 。",
    "点击 New upload。",
    "上传 release bundle，例如 CrustaVirus_DB_v1.0.0.zip。",
    "Title 填 CrustaVirus DB release v1.0.0。",
    "Resource type 选 Dataset 或 Software。",
    "填写真实 creators、ORCID、机构。",
    "Version 填 v1.0.0。",
    "选择 license。",
    "使用 Zenodo 的 reserve DOI 功能预留 DOI。",
    "把 DOI 写回 CITATION.cff、README、DATA_AVAILABILITY，再重新打包最终文件。",
    "最终检查无误后 Publish。",
]:
    num(step)

h1("8. 真实维护主体")
table(
    ["字段", "应该写", "不能写"],
    [
        ["Hosting institution", "真实学院/研究所/实验室", "TBD / upon acceptance"],
        ["Responsible PI", "姓名、职称、机构、邮箱", "CrustaVirus DB Team"],
        ["Technical maintainer", "姓名、邮箱、职责", "待分配"],
        ["Data curator", "姓名、邮箱、人工核对职责", "无"],
        ["Maintenance period", "2026-2031 至少 5 年", "看情况"],
        ["Funding/support", "项目号、机构运行经费、课题组经费", "planned only"],
        ["Domain owner", "单位/PI/课题组长期账号", "学生个人临时账号"],
        ["Server owner", "单位/课题组腾讯云账号", "个人临时云账号"],
    ],
)

h1("9. 投稿前最终检查")
h2("9.1 本地检查")
code("python release_gate.py\npython nar_readiness_check.py\npython tests\\run_all_tests.py\npython validate_database.py --check --report")
h2("9.2 公网检查")
for item in [
    "首页 HTTPS 可访问。",
    "搜索页、详情页、下载页可访问。",
    "API endpoint 可访问。",
    "所有下载文件 SHA256 与 manifest 一致。",
    "无需登录、无需注册。",
    "手机浏览器能打开。",
    "网站 footer 有备案号、license、contact、version、DOI。",
]:
    bullet(item)
h2("9.3 文档检查")
for item in [
    "PUBLIC_URL.txt 是真实 URL。",
    "CITATION.cff 无 TBD、无假 DOI。",
    "DATA_AVAILABILITY.md 有真实 URL、DOI、下载链接。",
    "SUSTAINABILITY.md 有真实 PI/机构/邮箱/5 年承诺。",
    "NOVELTY_COMPARISON.md 有引用、访问日期、对照方法。",
    "README.md 不再写 pre-release，除非确实还没正式发布。",
]:
    bullet(item)

h1("10. 官方依据")
for item in [
    "NAR Database Issue 指南：https://academic.oup.com/nar/pages/Ms_Prep_Database",
    "GitHub 新建仓库：https://docs.github.com/en/repositories/creating-and-managing-repositories/quickstart-for-repositories",
    "GitHub 大文件限制：https://docs.github.com/en/repositories/working-with-files/managing-large-files/about-large-files-on-github",
    "腾讯云备案流程：https://cloud.tencent.com/document/product/243/18909",
    "腾讯云备案云资源：https://cloud.tencent.com/document/product/243/18908/",
    "Zenodo DOI 预留：https://help.zenodo.org/docs/deposit/describe-records/reserve-doi/",
]:
    bullet(item)

h1("11. 你现在最应该做的 7 件事")
for step in [
    "确定最终维护主体：个人、PI、课题组还是学院。优先 PI/机构。",
    "确定域名：建议 crustavirusdb.org 或机构二级域名。",
    "建 GitHub private 仓库，先不要上传数据库和密钥。",
    "买腾讯云大陆服务器，若走大陆部署就立刻备案。",
    "备案等待期间继续人工核对 evidence、ICTV、host taxonomy。",
    "部署 staging 站点，只标注 pre-release。",
    "数据库冻结后再发布 DOI 和正式 v1.0.0。",
]:
    num(step)

p("最后的硬话：基础设施可以现在做，正式发布不能现在装。现在要抢时间启动备案和仓库，不要等数据库完美后才开始。")

doc.save(OUT_FILE)
print(OUT_FILE.resolve())
