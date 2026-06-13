#!/usr/bin/env python3
"""Generate NAR-formatted Word document from the paper draft."""
import re
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# --- Page setup ---
for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

# --- Default style ---
style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(12)
style.paragraph_format.line_spacing = 2.0  # Double-spaced for NAR
style.paragraph_format.space_after = Pt(0)
style.paragraph_format.space_before = Pt(0)

# Also set East Asian font fallback
rPr = style.element.get_or_add_rPr()
rFonts = OxmlElement('w:rFonts')
rFonts.set(qn('w:eastAsia'), '宋体')
rPr.append(rFonts)

# --- Helper functions ---
def add_heading_styled(text, level=1):
    """Add a heading with NAR formatting."""
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = 'Times New Roman'
        run.font.size = Pt(14) if level == 1 else Pt(12)
        if level == 1:
            run.font.bold = True
        rPr = run._element.get_or_add_rPr()
        rFonts = OxmlElement('w:rFonts')
        rFonts.set(qn('w:eastAsia'), '宋体')
        rPr.append(rFonts)
    h.paragraph_format.line_spacing = 2.0
    return h

def add_para(text, bold=False, italic=False, size=12, alignment=None, first_line_indent=None):
    """Add a paragraph with NAR formatting."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    rPr = run._element.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:eastAsia'), '宋体')
    rPr.append(rFonts)
    p.paragraph_format.line_spacing = 2.0
    if alignment:
        p.alignment = alignment
    if first_line_indent:
        p.paragraph_format.first_line_indent = Cm(first_line_indent)
    return p

def add_rich_para(segments):
    """Add paragraph with mixed formatting. segments = [(text, bold, italic), ...]"""
    p = doc.add_paragraph()
    for text, bold, italic in segments:
        run = p.add_run(text)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        run.bold = bold
        run.italic = italic
        rPr = run._element.get_or_add_rPr()
        rFonts = OxmlElement('w:rFonts')
        rFonts.set(qn('w:eastAsia'), '宋体')
        rPr.append(rFonts)
    p.paragraph_format.line_spacing = 2.0
    return p

def add_reference(text):
    """Add a reference entry."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(10)
    rPr = run._element.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:eastAsia'), '宋体')
    rPr.append(rFonts)
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.left_indent = Cm(1.27)
    p.paragraph_format.first_line_indent = Cm(-1.27)
    return p

# ============================================================
# TITLE PAGE
# ============================================================

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('AquaVir-KB: A Comprehensive Knowledge Base for Aquatic Invertebrate Viruses\nwith Genomic, Phylogenetic, and Functional Insights')
run.font.name = 'Times New Roman'
run.font.size = Pt(16)
run.bold = True
p.paragraph_format.line_spacing = 2.0
p.paragraph_format.space_after = Pt(24)

# Authors placeholder
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('[Author 1], [Author 2], [Author 3], [Author 4]*, [Author 5]')
run.font.name = 'Times New Roman'
run.font.size = Pt(12)
p.paragraph_format.line_spacing = 2.0

# Affiliations placeholder
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('[Department/School, University/Institute, City, Country]')
run.font.name = 'Times New Roman'
run.font.size = Pt(10)
run.italic = True
p.paragraph_format.line_spacing = 1.5

# Received date
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Received [Date]; Revised [Date]; Accepted [Date]')
run.font.name = 'Times New Roman'
run.font.size = Pt(10)
p.paragraph_format.space_before = Pt(24)

# ============================================================
# ABSTRACT
# ============================================================
add_heading_styled('ABSTRACT', level=1)

abstract_text = """Aquatic invertebrates constitute the productive foundation of global aquaculture, supplying over 28 million tonnes of food annually, yet viral diseases remain the largest source of economic loss in the sector. Despite the rapid accumulation of viral genomic data from metagenomic and metatranscriptomic surveys of marine and freshwater invertebrates, no centralized knowledge base provides systematically curated, evidence-graded virus–host associations for this ecologically and economically critical group. Here we present AquaVir-KB, a comprehensive, evidence-driven knowledge base dedicated to viruses infecting aquatic invertebrates. The database integrates 934 virus species spanning seven phyla (Arthropoda, Mollusca, Echinodermata, Cnidaria, Porifera, Annelida, and Rotifera), 11,884 viral isolates, 26,894 proteins, and 331,030 structured literature evidence records derived from 8,511 peer-reviewed publications (98.0% PMID coverage, 95.4% DOI coverage). Built on a seven-layer, 116-table relational architecture, AquaVir-KB uniquely combines ICTV-aligned taxonomy, NCBI sequence integration, genome annotation, protein functional domain classification (56.3% of 26,894 proteins annotated), phylogenetic analysis (131 RdRP sequences classified across 18 families), and geographic distribution data (48 countries, 7 continents) within a single resource. The database introduces a three-tier evidence grading system that explicitly labels the evidentiary basis of each virus–host association. AquaVir-KB is freely accessible via a REST API with full-text search, provides bulk data downloads (117.6 MB, CC-BY 4.0), and supports Docker-based local deployment."""
add_para(abstract_text, size=12)

# Database URL
p = doc.add_paragraph()
run = p.add_run('Database URL: ')
run.font.name = 'Times New Roman'
run.font.size = Pt(12)
run.bold = True
run = p.add_run('[TBD upon deployment at https://aquavir-kb.org]')
run.font.name = 'Times New Roman'
run.font.size = Pt(12)
p.paragraph_format.line_spacing = 2.0

# ============================================================
# INTRODUCTION
# ============================================================
add_heading_styled('INTRODUCTION', level=1)

intro_para1 = """Aquatic invertebrates form the productive backbone of global aquaculture—the fastest-growing food production sector worldwide—contributing 28.9 million tonnes annually, comprising 11.2 million tonnes of crustaceans valued at approximately USD 78 billion and 17.7 million tonnes of mollusks [1]. Viral diseases are the single largest source of economic loss in invertebrate aquaculture. White spot syndrome virus (WSSV; family Nimaviridae) has caused cumulative global losses estimated at USD 15–20 billion since its emergence in the early 1990s, devastating shrimp farming operations across Asia and the Americas [2, 3]. Taura syndrome virus (TSV; Dicistroviridae) emerged in Ecuador in 1992 and spread throughout the Americas within a decade, causing regional production collapses with mortality rates exceeding 90% in susceptible stocks [5]. Infectious hypodermal and hematopoietic necrosis virus (IHHNV; Parvoviridae) causes growth retardation and cuticular deformities in Penaeus vannamei, resulting in substantial economic losses through reduced harvest weight rather than direct mortality [6]. Yellow head virus (YHV; Roniviridae) has caused episodic mass mortality events in Penaeus monodon culture across Southeast Asia since 1990 [3]. In mollusks, ostreid herpesvirus 1 (OsHV-1; Malacoherpesviridae) and its microvariant genotype OsHV-1 μVar have driven mass mortality events in Pacific oyster (Crassostrea gigas) hatcheries and grow-out operations across France, the United Kingdom, Ireland, Australia, and New Zealand since 2008, with mortality rates frequently reaching 80–100% in spat and juveniles [8, 7]. Haliotid herpesvirus 1 (HaHV-1) causes acute ganglioneuritis in farmed and wild abalone (Haliotis spp.), with outbreaks documented in Taiwan and Australia [9]."""
add_para(intro_para1, first_line_indent=1.27)

intro_para2 = """Beyond the aquaculture context, aquatic invertebrates harbor an immense and largely uncharacterized viral diversity. Metagenomic and metatranscriptomic surveys have revealed that marine invertebrates host thousands of novel RNA and DNA viruses spanning all major viral genome types and representing numerous previously undescribed lineages [10, 11]. Coral holobionts alone are estimated to contain thousands of uncharacterized viral populations that may play critical roles in coral health and disease [15]. Sponges—among the most ancient metazoans—host diverse viromes that may contribute to their remarkable chemical defense systems [10]. Echinoderms, including commercially valuable sea cucumbers (Apostichopus japonicus) subject to intensive aquaculture in China, harbor novel picornaviruses and parvoviruses of unknown pathogenic potential. Despite this dual economic and biological significance, a centralized, systematically curated knowledge base for aquatic invertebrate viruses has been conspicuously absent."""
add_para(intro_para2, first_line_indent=1.27)

intro_para3 = """Several publicly available databases serve adjacent functions but do not address the specific need for curated, evidence-graded virus–host association data for aquatic invertebrates. NCBI Virus provides the broadest sequence coverage with links to GenBank records and metadata, but it does not offer structured evidence grading, curated host verification, or literature traceability beyond the associated GenBank reference. The International Committee on Taxonomy of Viruses (ICTV) Virus Metadata Resource (VMR) provides authoritative taxonomic classification for all recognized virus species [27, 28], yet its host metadata are limited to coarse categorical annotations that do not resolve aquatic from terrestrial invertebrate hosts. Virus-Host DB [16] computationally infers virus–host associations from NCBI Taxonomy fields, but these computationally derived associations lack primary literature verification and offer no evidence quality stratification. The Reference Viral Database (RVDB) [17] clusters RNA virus sequences at nucleotide and protein levels but provides no curated host associations. IMG/VR v3 [18] serves the viral metagenomics community with environmental virus sequence catalogs (2.3 million uncultivated viral genomes) but lacks structured host metadata beyond environmental provenance. The ViPTree server [19] provides phylogenetic placement services for novel viral sequences but does not curate host information."""
add_para(intro_para3, first_line_indent=1.27)

intro_para4 = """Here we describe AquaVir-KB, a comprehensive, evidence-driven knowledge base designed to fill this gap. The database was constructed through systematic integration of six primary data sources: (i) the ICTV Virus Metadata Resource (MSL41, 2025 release); (ii) NCBI GenBank, RefSeq, and Nucleotide databases; (iii) Europe PMC and PubMed literature databases; (iv) InterPro, Pfam, KEGG, and UniProt protein annotation resources; (v) the NCBI Sequence Read Archive (SRA) for metagenomic data; and (vi) GBIF and OBIS for biogeographic host occurrence data. These data were organized into a seven-layer relational architecture designed to separate concerns while maintaining referential integrity through foreign key constraints. A central methodological contribution is the explicit three-tier evidence grading system (high/medium/low) applied to every virus–host association. Below, we describe the database architecture, data integration methodology, quality control procedures, and the content of the first public release."""
add_para(intro_para4, first_line_indent=1.27)

# ============================================================
# MATERIALS AND METHODS
# ============================================================
add_heading_styled('MATERIALS AND METHODS', level=1)

# Section 1
add_heading_styled('1. Data Collection, Verification, and Management', level=2)

methods_1a = """ICTV Taxonomy Integration. The complete ICTV Master Species List (MSL41, released March 2025) was downloaded from the ICTV website and imported into the ictv_taxonomy and ictv_vmr tables. The ICTV database, maintained as the authoritative reference for virus taxonomy [27, 28], provides the taxonomic framework to which all AquaVir-KB virus entries are mapped. This yielded 17,554 taxonomy records spanning all ICTV-recognized taxonomic ranks (realm through species) and 19,271 virus metadata records. A systematic cross-referencing procedure matched ICTV species names against the existing virus_master table using exact and normalized match strategies, establishing 1,732 virus–ICTV mappings. A gap analysis identified 40 ICTV-listed aquatic invertebrate-associated viruses absent from the initial database build—predominantly 22 Aquambidensovirus asteroid1–22 (Parvoviridae) from Echinodermata, 11 crustacean and mollusk parvoviruses and iridoviruses, and 7 mollusk-associated ourmiaviruses and malacoherpesviruses."""
add_para(methods_1a, first_line_indent=1.27)

methods_1b = """NCBI GenBank and RefSeq Integration. Viral nucleotide accessions were retrieved from NCBI using Entrez E-utilities. Query strategies employed both targeted organism-level searches and broad keyword-based searches across host common names (oyster, mussel, abalone, clam, scallop, coral, jellyfish, sea cucumber, starfish, sea urchin, sponge, polychaete, rotifer). Searches were executed against the NCBI Nucleotide database with date filters to capture both historical and recent (2023–2026) accessions. Results were batch-validated using NCBI ESummary, extracting AccessionVersion, Organism, Title, Length, and associated BioProject/BioSample identifiers. Redundant accessions were deduplicated at the accession level. The final dataset comprises 11,884 viral isolates, of which 7,842 (66.0%) have associated nucleotide sequence data, 2,587 possess genome-level accession numbers, and 6,262 (52.7%) are linked to specific literature references."""
add_para(methods_1b, first_line_indent=1.27)

methods_1c = """Literature Collection and Structured Evidence Extraction. Systematic literature searches were executed against Europe PMC and PubMed using 16 query strategies covering all seven target phyla (Supplementary Table S2). Retrieved publications were deduplicated by PMID (primary key) and DOI (secondary key). After deduplication, 8,511 unique references were retained. Structured evidence extraction employed a multi-stage, progressively deepening pipeline: Stage 1 applied automated keyword and regular expression pattern matching against titles and abstracts; Stage 2 processed full-text XML for 4,371 downloadable articles (44.6% of attempted downloads); Stage 3 implemented a fuzzy token-matching algorithm against controlled virus name and host name vocabularies. The complete pipeline produced 331,030 structured evidence records."""
add_para(methods_1c, first_line_indent=1.27)

methods_1d = """Quality Control and Data Hygiene. Duplicate virus entries (n = 9 groups) arising from case-variant naming during multi-source import were identified through case-insensitive canonical name matching and consolidated, migrating all associated records to the canonical entry. Non-target records (n = 380; viruses of algae, vertebrates, fungi, terrestrial plants, and non-aquatic organisms) inadvertently incorporated during broad taxonomic import sweeps were flagged and excluded from all public-facing views. A stratified random sampling validation of auto-imported evidence records (n = 300; 100 per evidence strength tier) assessed seven format and completeness criteria, yielding an overall precision of 89.9%."""
add_para(methods_1d, first_line_indent=1.27)

# Section 2
add_heading_styled('2. Sequence Data Acquisition and Genome Processing', level=2)

methods_2 = """Viral nucleotide sequences were retrieved from NCBI Nucleotide and RefSeq for all 7,842 isolates with sequence data. Genome metadata—including genome length, GC content, molecule type, sequence completeness, and associated BioProject/BioSample identifiers—were extracted from GenBank flat file annotations via EFetch. Genome length quality was assessed against expected ranges for each viral family based on ICTV reference data; outliers (>3 standard deviations from the family median) were flagged for manual review. Protein-coding sequences (CDS) were extracted from GenBank feature tables, yielding 26,894 viral protein accessions with validated amino acid translations (mean length: 373 aa, range: 19–7,195 aa). For 3,511 isolates lacking curated CDS annotations in GenBank—predominantly recently deposited metagenomic assemblies—open reading frame (ORF) prediction was performed using Prodigal v2.6.3 [25] in metagenomic mode with a minimum ORF length of 150 nucleotides. This generated 61,339 reannotated ORFs."""
add_para(methods_2, first_line_indent=1.27)

# Section 3
add_heading_styled('3. Genome Annotation, Non-redundant Protein Database, and Core Gene Identification', level=2)

methods_3 = """Non-redundant Protein Clustering. To reduce sequence redundancy across the 26,894 viral proteins while preserving functional diversity, proteins were clustered using CD-HIT v4.8.1 [26] at a 95% amino acid identity threshold. This produced 16,730 non-redundant (NR) protein clusters, with a cluster size distribution ranging from singletons (n = 12,847 clusters) to large multi-member families (largest cluster: 847 members). Core Gene Identification. Core viral genes were identified through profile Hidden Markov Model (HMM) searches against the Pfam v36.0 database [49] and NCBI Conserved Domain Database (CDD) [45]. The Pfam HMM library was searched using hmmscan (HMMER v3.3.2) [24] with an E-value reporting threshold of 1 × 10⁻⁵. A total of 3,642 proteins met these criteria and were classified as core viral genes. Genome type classification followed ICTV genome composition standards."""
add_para(methods_3, first_line_indent=1.27)

# Section 4
add_heading_styled('4. Protein Functional Annotation and Domain Architecture Analysis', level=2)

methods_4 = """Conserved Domain Annotation. All 26,894 viral proteins were submitted to the NCBI CDD [45] batch search service, yielding 65,943 domain assignments across 52,511 proteins from NCBI CDD, 5,809 from protein name inference, 3,871 from UniProt keyword mapping [46], and 3,752 from rule-based classification. Functional Category Assignment. A rules-based inference engine was developed to map domain names and descriptions to six functional categories (RdRP, replication, structural, host_interaction, metabolism, assembly) using a hierarchical set of 50+ regular expression patterns. A secondary inference layer processed protein names and gene symbols for proteins lacking domain matches but carrying informative annotations. After the full pipeline execution, 15,142 proteins (56.3%) received functional category assignments: 5,809 RdRP (21.6%), 3,960 structural (14.7%), 3,904 replication (14.5%), 1,090 metabolism (4.1%), 365 host_interaction (1.4%), and 14 assembly (0.1%). The remaining 11,752 proteins (43.7%) represent uncharacterized ORFs. InterPro, KEGG, and Gene Ontology Integration. Protein sequences were queried against InterPro v97.0 [47] via the InterPro REST API, yielding domain annotations for 427 proteins with associated GO term mappings (3,452 annotations across 176 unique GO identifiers). KEGG Orthology assignments were retrieved for 2,814 proteins via the KEGG REST API [48], with 4,294 protein-pathway links across 924 unique KEGG pathways. Three-dimensional structural models were predicted for 52 representative proteins using ESMFold [58] via the ESMFold API."""
add_para(methods_4, first_line_indent=1.27)

# Section 5
add_heading_styled('5. Phylogenetic and Evolutionary Relationship Analysis', level=2)

methods_5 = """RdRp Sequence Collection and Alignment. RNA-dependent RNA polymerase sequences constitute the most conserved and phylogenetically informative gene across RNA viruses. For 131 representative RdRp sequences from unclassified or ambiguously classified viruses, phylogenetic analysis was performed. Multiple sequence alignment was performed using MAFFT v7.520 [20] with the L-INS-i algorithm, selected for its accuracy on highly divergent sequences. Alignment quality was assessed using trimAl v1.4 [50] with the -automated1 heuristic. Maximum-Likelihood Phylogenetic Inference. Maximum-likelihood trees were constructed using IQ-TREE v2.2.0 [21] with automatic model selection via ModelFinder [22], assessing 286 DNA models or 546 protein models and selecting the best-fit model according to Bayesian Information Criterion (BIC). Branch support was evaluated using 1,000 ultrafast bootstrap replicates [21] and the Shimodaira-Hasegawa approximate likelihood ratio test (SH-aLRT) [23]. Family-level classification was assigned based on monophyletic clustering (bootstrap support ≥70% and SH-aLRT ≥80%) with known reference taxa, yielding family-level assignments across 18 viral families."""
add_para(methods_5, first_line_indent=1.27)

# Section 6
add_heading_styled('6. Applied Knowledge Base Construction', level=2)

methods_6 = """AquaVir-KB is built on a seven-layer, 116-table relational data model (Figure 1). Layer 1 (Core Virus) contains virus taxonomy, genome type, discovery context, and ICTV mappings. Layer 2 (Host) catalogs aquatic invertebrate host species with associated taxonomy profiles. Layer 3 (Evidence) contains structured literature-derived evidence records with grading. Layer 4 (Literature) stores bibliographic metadata with full-text source tracking. Layer 5 (Protein) integrates functional annotations from NCBI CDD, InterPro, KEGG, and UniProt. Layer 6 (Geography/Ecology) records sample collection metadata, GBIF/OBIS occurrence records, and temperature/virulence profiles. Layer 7 (Curation) implements provenance tracking, curation audit logging, and data conflict management. The database's core methodological innovation is the explicit three-tier evidence grading system embedded in the evidence_records table. Each of the 331,030 evidence records carries an evidence_strength field (high/medium/low) and an extraction_method field documenting the data provenance pathway. All data modifications are recorded in curation_logs, and all records carry provenance annotations linking to their source databases. Data conflicts are recorded in curation_conflicts with severity grading and resolution status tracking."""
add_para(methods_6, first_line_indent=1.27)

# Section 7
add_heading_styled('7. Database Deployment and Web Interface', level=2)

methods_7 = """A REST API was implemented using FastAPI (Python 3.12) [64] providing programmatic access to all public database views. Full-text search across virus names, host names, taxonomy fields, and geographic locations is implemented via SQLite FTS5. API responses are returned in JSON format with pagination and CORS headers. Interactive API documentation is generated via Swagger UI and ReDoc. The complete application stack is containerized via Docker Compose, comprising PostgreSQL 16, FastAPI with Uvicorn (4 workers), and Nginx reverse proxy with rate limiting (10 requests/second). Deployment requires a single command after setting the database password environment variable. Complete database snapshots are provided as five TSV files (117.6 MB total) under CC-BY 4.0 license. The database will be deposited in Zenodo with a versioned DOI upon publication. The design of AquaVir-KB follows the FAIR Guiding Principles [55], ensuring Findability (DOI, rich metadata), Accessibility (REST API, bulk download), Interoperability (ICTV taxonomy, NCBI Taxonomy, GO terms, KEGG pathways), and Reusability (CC-BY 4.0, provenance tracking, Docker reproducibility)."""
add_para(methods_7, first_line_indent=1.27)

# ============================================================
# RESULTS
# ============================================================
add_heading_styled('RESULTS', level=1)

add_heading_styled('1. Database Overview', level=2)

results_1 = """AquaVir-KB v1.0 contains 934 virus species (public release) across seven phyla, infecting 183 aquatic invertebrate host species (Table 1). The database integrates 11,884 viral isolates, 26,894 proteins, 331,030 structured literature evidence records, and 8,511 literature references spanning 1950–2026. Data provenance is tracked through 100,599 source attribution records. The total curated data volume is 624 MB (SQLite). The Arthropoda dominance (67.0% of virus species) reflects over four decades of intensive research on penaeid shrimp viral pathogens. Mollusca (22.2%) represent the second-largest group, while the remaining phyla collectively represent 10.8% of virus records, reflecting the nascent state of virological investigation in these taxa."""
add_para(results_1, first_line_indent=1.27)

# Table 1
add_heading_styled('Table 1. Virus and host distribution across aquatic invertebrate phyla.', level=3)
table1 = doc.add_table(rows=11, cols=6)
table1.style = 'Table Grid'
headers1 = ['Phylum', 'Virus Species', '%', 'Host Species', '%', 'Representative Host Genera']
for i, h in enumerate(headers1):
    cell = table1.rows[0].cells[i]
    cell.text = h
    for p in cell.paragraphs:
        for run in p.runs:
            run.font.size = Pt(9)
            run.font.bold = True
            run.font.name = 'Times New Roman'

data1 = [
    ['Arthropoda', '626', '67.0', '111', '60.7', 'Penaeus, Litopenaeus, Macrobrachium, Scylla'],
    ['Mollusca', '207', '22.2', '43', '23.5', 'Crassostrea, Mytilus, Haliotis, Ruditapes'],
    ['Echinodermata', '37', '4.0', '10', '5.5', 'Apostichopus, Strongylocentrotus, Asterias'],
    ['Cnidaria', '18', '1.9', '10', '5.5', 'Acropora, Nematostella, Hydra, Aurelia'],
    ['Porifera', '13', '1.4', '5', '2.7', 'Amphimedon, Stylissa, Aplysina'],
    ['Annelida', '3', '0.3', '2', '1.1', 'Arenicola, Perinereis'],
    ['Rotifera', '1', '0.1', '0', '0.0', '—'],
    ['Nematoda', '1', '0.1', '1', '0.5', 'Caenorhabditis'],
    ['Platyhelminthes', '1', '0.1', '1', '0.5', 'Schmidtea'],
    ['Multiple phyla', '27', '2.9', '—', '—', 'Cross-phylum associations'],
]
for i, row_data in enumerate(data1):
    for j, val in enumerate(row_data):
        cell = table1.rows[i+1].cells[j]
        cell.text = val
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.size = Pt(9)
                run.font.name = 'Times New Roman'

# Results sections 2-7 (condensed for Word)
add_heading_styled('2. Taxonomic and Genomic Coverage', level=2)
add_para("""Family-level classification covers 92.7% of public viruses (866/934). The database spans 63 unique viral families and 53 genera. The 20 most abundant families are dominated by Picornavirales (n = 189), Picornaviridae (n = 150), and Nucleocytoviricota (n = 90). Single-stranded positive-sense RNA genomes dominate, constituting 65.9% of species (n = 615), consistent with the prevalence of picornavirus-like and marnavirus-like genomes in global marine RNA virome surveys [11, 13, 14]. Double-stranded DNA viruses (n = 155, 16.6%) represent the second-largest group. Discovery context analysis reveals that 96.4% of viruses were identified through metagenomic or metatranscriptomic approaches, with only 4 viruses (0.4%) isolated and cultured. The Serratus project's petabase-scale alignment of 5.7 million SRA samples [12] and the Tara Oceans expedition's characterization of marine RNA virus diversity [13, 14] represent the current frontier of aquatic virus discovery.""", first_line_indent=1.27)

add_heading_styled('3. Genome and Proteome Features', level=2)
add_para("""Genome lengths among the 7,842 isolates with sequence data range from 16 bp to over 1.2 Mb, with a mean of 120.4 kb and a median of 9.2 kb. The viral proteome of 26,894 proteins spans 19 to 7,195 amino acids (mean: 373 aa). Functional category assignment reveals a proteome dominated by replication-associated functions: RNA-dependent RNA polymerase (5,809 proteins, 21.6%), structural proteins (3,960, 14.7%), replication-associated non-RdRP proteins (3,904, 14.5%), metabolism-associated proteins (1,090, 4.1%), host-interaction proteins (365, 1.4%), and assembly-related proteins (14, 0.1%). Non-redundant clustering reduced the protein set to 16,730 NR clusters. Core gene identification via Pfam HMM profiling identified 3,642 proteins belonging to essential viral functional classes.""", first_line_indent=1.27)

add_heading_styled('4. Literature Evidence and Virus–Host Association Records', level=2)
add_para("""The database contains 331,030 structured evidence records derived from 8,511 publications spanning 1950–2026. Literature traceability is high: 98.0% of references carry PubMed IDs and 95.4% carry DOIs. The evidence strength stratification reveals a pyramid-shaped distribution: 4,906 records (1.5%) classified as high-grade, 159,522 (48.2%) as medium-grade, and 166,602 (50.3%) as low-grade. For the nine most intensively studied virus–host model systems, evidence quality is substantially elevated above the database-wide average. WSSV has 17.0% high-grade evidence among 10,550 associated records, reflecting over three decades of experimental infection research including early fulfillment of Koch's postulates [4]. YHV (7.0% high-grade among 2,008 records) and IHHNV (7.1% among 4,159 records) similarly benefit from decades of experimental characterization. Comprehensive reviews of shrimp viral diseases in Asia [30] and the Americas [31], mollusk herpesvirus epidemiology [36, 37, 38], abalone ganglioneuritis pathobiology [39, 41], and emerging nodavirus infections in echinoderms [40] provide the primary literature foundation. Among 9,548 infection records, 842 (8.8%) represent experimentally confirmed infections, 4 represent pathology observations, and 2,271 (23.8%) represent disease outbreak associations. The most virus-rich individual host species include Procambarus clarkii (250 associated viruses), Crustacea (unspecified, 205 viruses), Litopenaeus vannamei (189 viruses), and Acanthaster planci (154 viruses).""", first_line_indent=1.27)

add_heading_styled('5. Virus Isolate and Protein Detail Records', level=2)
add_para("""The isolate detail layer provides per-accession granularity for each virus species. For the 7,842 isolates with sequence data, protein-coding sequences are organized with individual protein accessions, genomic coordinates, amino acid sequences, and functional annotations. The protein detail layer integrates domain architecture from NCBI CDD (65,943 assignments), InterPro domain classifications, GO term mappings (3,452 annotations across 176 unique GO identifiers), KEGG pathway links (4,294 protein-pathway associations across 924 pathways), and UniProt cross-references (11,351 proteins). Three-dimensional structural models are available for 52 representative proteins.""", first_line_indent=1.27)

add_heading_styled('6. Phylogenetic and Evolutionary Relationships', level=2)
add_para("""Phylogenetic analysis of 131 representative RdRp sequences from unclassified or ambiguously classified RNA viruses yielded family-level classifications across 18 viral families, with robust support (bootstrap ≥70%, SH-aLRT ≥80%). Sixty-eight viruses (7.3% of 934 public species) currently lack family-level classification and represent candidates for systematic phylogenetic placement in the ongoing classification pipeline. These taxonomic orphans disproportionately derive from Porifera (46% unclassified), Cnidaria (28% unclassified), and Echinodermata (22% unclassified)—phyla for which reference virological data are sparse.""", first_line_indent=1.27)

add_heading_styled('7. Database Access and User Interface', level=2)
add_para("""The REST API supports parameterized queries across all public database views, with response pagination and JSON output. Full-text search enables rapid retrieval across virus names, host names, taxonomy, and geographic fields. Five TSV files (117.6 MB) provide the complete public dataset. Sample collection metadata span 48 countries across seven continents, with the strongest coverage in Asia (1,329 records), North America (428), and Oceania (64). Metadata for 2,598 SRA runs from aquatic invertebrate metagenomic and metatranscriptomic studies are indexed, of which 400 have detectable viral sequences via NCBI Taxonomy Analysis.""", first_line_indent=1.27)

# ============================================================
# DISCUSSION
# ============================================================
add_heading_styled('DISCUSSION', level=1)

add_para("""AquaVir-KB fills a long-standing gap in virology informatics by providing the first systematically curated, evidence-graded knowledge base dedicated to aquatic invertebrate viruses. The database's primary contribution is methodological: by making explicit the evidentiary basis of each virus–host association through a transparent, three-tier grading system, AquaVir-KB transforms what is currently implicit in the literature—and entirely absent from sequence databases—into structured, queryable metadata. For a field in which 96.4% of known viruses have been discovered through sequencing rather than classical isolation, the distinction between a Koch's postulate-confirmed pathogen and a metagenomic co-occurrence has immediate practical consequences for biosecurity policy, aquaculture disease management, and research prioritization.""", first_line_indent=1.27)

add_para("""AquaVir-KB complements rather than duplicates existing virus databases. It serves a fundamentally different purpose from sequence repositories (NCBI Virus, GenBank), taxonomic authorities (ICTV VMR), and computational inference platforms (Virus-Host DB). The database's unique value proposition lies in the combination of (i) explicit evidence grading, (ii) per-record literature provenance (98.0% PMID, 95.4% DOI), (iii) phylum-specific scope for aquatic invertebrates, and (iv) multi-dimensional data integration within a single relational framework. The 2025 NAR Database Issue collection [54] documents the continued growth of molecular biology databases, within which AquaVir-KB occupies a previously unfilled niche. The database's functional annotation pipeline leverages Pfam [49] for core gene identification and ultrafast bootstrap approximation [51] for phylogenetic support assessment. Geographic host occurrence data were integrated from GBIF [56] and OBIS [57]. Future releases will incorporate automated virus discovery tools including VirSorter2 [61], Cenote-Taker 2 [60], and VIBRANT [52] for enhanced metagenomic mining. The expansion of known ssRNA phage genomes [59] and the unprecedented genomic diversity revealed by arthropod RNA virus surveys [65] underscore the discovery potential remaining in under-sampled aquatic invertebrate taxa.""", first_line_indent=1.27)

add_para("""Several limitations of the current release warrant transparent acknowledgment. First, phylum coverage is highly uneven, with 89.2% of virus species concentrated in Arthropoda (67.0%) and Mollusca (22.2%). This reflects the distribution of research effort rather than the true diversity of aquatic invertebrate viruses. Coral reef cnidarians, deep-sea echinoderms, polar sponges, and freshwater annelids are grossly under-sampled. Recent metagenomic surveys of sponge-associated viral communities [43] and coral holobiont viromes [42] suggest substantial undiscovered viral diversity. Second, the evidence pyramid is bottom-heavy: only 1.5% of evidence records achieve the high grade. While this accurately reflects the metagenomic dominance of virus discovery, it means most entries await experimental validation. Third, full-text access remains incomplete (44.6% download success rate). Fourth, Chinese-language aquaculture virology literature—published in CNKI-indexed journals—is underrepresented. Fifth, the production web deployment (public URL, HTTPS, Zenodo DOI) remains pending at manuscript preparation.""", first_line_indent=1.27)

add_para("""AquaVir-KB is designed as an evolving resource. Several development trajectories are planned: (i) systematic RdRP HMM profiling against assembled contigs from the 2,598 indexed SRA runs, modeled on the Serratus approach [12], with conservative estimates suggesting 500–1,500 additional virus species; (ii) Chinese literature integration via CNKI-indexed journal text mining; (iii) systematic phylogenetic placement of 68 unclassified viruses; (iv) transition from static release to automated weekly updates from NCBI GenBank, Europe PMC, and NCBI SRA; (v) community curation interface enabling domain expert review and augmentation of evidence records; and (vi) targeted data collection for currently unrepresented aquatic invertebrate phyla, aligned with emerging frameworks for aquatic food safety risk assessment [62, 63].""", first_line_indent=1.27)

# ============================================================
# DATA AVAILABILITY
# ============================================================
add_heading_styled('DATA AVAILABILITY', level=1)
add_para("""AquaVir-KB is freely accessible at [URL TBD]. The complete database (934 virus species, 183 hosts, 10,430 isolates, 331,030 evidence records, 8,511 references) is available for download as TSV files (117.6 MB) under the CC-BY 4.0 license. Source code for the data processing pipeline, database schema DDL, and Docker deployment configuration is available at [repository URL TBD]. A Docker image supporting single-command local deployment is provided. The database will be deposited in Zenodo with a versioned DOI upon publication. Non-target records (n = 380) are excluded from the public release with documented exclusion criteria.""", first_line_indent=1.27)

# ============================================================
# SUPPLEMENTARY INFORMATION
# ============================================================
add_heading_styled('SUPPLEMENTARY INFORMATION', level=1)

add_heading_styled('Supplementary Table S1. Feature comparison of AquaVir-KB with existing virus databases.', level=2)
add_para("""[See Supplementary Material for the complete 7-database, 15-feature comparison matrix.]""")

add_heading_styled('Supplementary Table S2. Literature search strategies and result counts.', level=2)
add_para("""[See Supplementary Material for the complete 16-search strategy table with hit counts.]""")

# ============================================================
# ACKNOWLEDGEMENTS
# ============================================================
add_heading_styled('ACKNOWLEDGEMENTS', level=1)
add_para('[To be added]')

add_heading_styled('FUNDING', level=1)
add_para('[To be added]')

# ============================================================
# REFERENCES
# ============================================================
add_heading_styled('REFERENCES', level=1)

references = [
    "1. FAO. (2024) The State of World Fisheries and Aquaculture 2024 (SOFIA). Food and Agriculture Organization of the United Nations, Rome. https://www.fao.org/documents/card/en/c/cc0683en",
    "2. Stentiford, G.D., Neil, D.M., Peeler, E.J., Shields, J.D., Small, H.J., Flegel, T.W., Vlak, J.M., Jones, B., Morado, F., Moss, S., Lotz, J., Bartholomay, L., Behringer, D.C., Hauton, C. and Lightner, D.V. (2012) Disease will limit future food supply from the global crustacean fishery and aquaculture sectors. J. Invertebr. Pathol., 110(2), 141–157. DOI: 10.1016/j.jip.2012.03.013",
    "3. Flegel, T.W. (2012) Historic emergence, impact and current status of shrimp pathogens in Asia. J. Invertebr. Pathol., 110(2), 166–173. DOI: 10.1016/j.jip.2012.03.004",
    "4. Lightner, D.V. and Redman, R.M. (1998) Shrimp diseases and current diagnostic methods. Aquaculture, 164(1–4), 201–220. DOI: 10.1016/S0044-8486(98)00187-2",
    "5. Lightner, D.V., Redman, R.M., Pantoja, C.R., Tang, K.F.J., Noble, B.L., Schofield, P., Mohney, L.L., Nunan, L.M. and Navarro, S.A. (2012) Historic emergence, impact and current status of shrimp pathogens in the Americas. J. Invertebr. Pathol., 110(2), 174–183. DOI: 10.1016/j.jip.2012.03.006",
    "6. Rai, P., Safeena, M.P., Krabsetsve, K., La Fauce, K., Owens, L. and Karunasagar, I. (2012) Genomics, molecular epidemiology and diagnostics of infectious hypodermal and hematopoietic necrosis virus. Indian J. Virol., 23(2), 203–214. DOI: 10.1007/s13337-012-0070-7",
    "7. Segarra, A., Pépin, J.-F., Arzul, I., Morga, B., Faury, N. and Renault, T. (2010) Detection and description of a particular Ostreid herpesvirus 1 genotype associated with massive mortality outbreaks of Pacific oysters, Crassostrea gigas, in France in 2008. Virus Res., 153(1), 92–99. DOI: 10.1016/j.virusres.2010.07.011",
    "8. Pernet, F., Lupo, C., Bacher, C. and Whittington, R.J. (2016) Infectious diseases in oyster aquaculture require a new integrated approach. Phil. Trans. R. Soc. B, 371(1689), 20150213. DOI: 10.1098/rstb.2015.0213",
    "9. Hooper, C., Hardy-Smith, P. and Handlinger, J. (2007) Ganglioneuritis causing high mortalities in farmed Australian abalone (Haliotis laevigata and Haliotis rubra). Aust. Vet. J., 85(5), 188–193. DOI: 10.1111/j.1751-0813.2007.00155.x",
    "10. Shi, M., Lin, X.-D., Tian, J.-H., Chen, L.-J., Chen, X., Li, C.-X., Qin, X.-C., Li, J., Cao, J.-P., Eden, J.-S., Buchmann, J., Wang, W., Xu, J., Holmes, E.C. and Zhang, Y.-Z. (2016) Redefining the invertebrate RNA virosphere. Nature, 540(7634), 539–543. DOI: 10.1038/nature20167",
    "11. Wolf, Y.I., Silas, S., Wang, Y., Wu, S., Bocek, M.J., Kazlauskas, D., Krupovic, M., Fire, A., Dolja, V.V. and Koonin, E.V. (2020) Doubling of the known set of RNA viruses by metagenomic analysis of an aquatic virome. Nat. Microbiol., 5, 1262–1270. DOI: 10.1038/s41564-020-0755-4",
    "12. Edgar, R.C., Taylor, J., Lin, V., Altman, T., Barbera, P., Meleshko, D., Lohr, D., Novakovsky, G., Buchfink, B., Al-Shayeb, B., Banfield, J.F., de la Peña, M., Korobeynikov, A., Chikhi, R. and Babaian, A. (2022) Petabase-scale sequence alignment catalyses viral discovery. Nature, 602(7895), 142–147. DOI: 10.1038/s41586-021-04332-2",
    "13. Dominguez-Huerta, G., Zayed, A.A., Wainaina, J.M., Guo, J., Tian, F., Pratama, A.A., Bolduc, B., Mohssen, M., Zablocki, O., Pelletier, E., ... and Sullivan, M.B. (2022) Diversity and ecological footprint of global ocean RNA viruses. Science, 376(6598), 1202–1208. DOI: 10.1126/science.abn6358",
    "14. Zayed, A.A., Wainaina, J.M., Dominguez-Huerta, G., Pelletier, E., Guo, J., Mohssen, M., Tian, F., Pratama, A.A., Bolduc, B., Zablocki, O., ... and Sullivan, M.B. (2022) Cryptic and abundant marine viruses at the evolutionary origins of Earth's RNA virome. Science, 376(6589), 156–162. DOI: 10.1126/science.abm5847",
    "15. Vega Thurber, R., Payet, J.P., Thurber, A.R. and Correa, A.M.S. (2017) Virus–host interactions and their roles in coral reef health and disease. Nat. Rev. Microbiol., 15(4), 205–216. DOI: 10.1038/nrmicro.2016.176",
    "16. Mihara, T., Nishimura, Y., Shimizu, Y., Nishiyama, H., Yoshikawa, G., Uehara, H., Hingamp, P., Goto, S. and Ogata, H. (2016) Linking virus genomes with host taxonomy. Viruses, 8(3), 66. DOI: 10.3390/v8030066",
    "17. Goodacre, N., Aljanahi, A., Nandakumar, S., Mikailov, M. and Khan, A.S. (2018) A Reference Viral Database (RVDB) to enhance bioinformatics analysis of high-throughput sequencing for novel virus detection. mSphere, 3(2), e00069-18. DOI: 10.1128/mSphereDirect.00069-18",
    "18. Roux, S., Páez-Espino, D., Chen, I.-M.A., Palaniappan, K., Ratner, A., Chu, K., Reddy, T.B.K., Nayfach, S., Schulz, F., Call, L., Neches, R.Y., Woyke, T., Ivanova, N.N., Eloe-Fadrosh, E.A. and Kyrpides, N.C. (2021) IMG/VR v3: an integrated ecological and evolutionary framework for interrogating genomes of uncultivated viruses. Nucleic Acids Res., 49(D1), D764–D775. DOI: 10.1093/nar/gkaa946",
    "19. Nishimura, Y., Yoshida, T., Kuronishi, M., Uehara, H., Ogata, H. and Goto, S. (2017) ViPTree: the viral proteomic tree server. Bioinformatics, 33(15), 2379–2380. DOI: 10.1093/bioinformatics/btx157",
    "20. Katoh, K. and Standley, D.M. (2013) MAFFT multiple sequence alignment software version 7: improvements in performance and usability. Mol. Biol. Evol., 30(4), 772–780. DOI: 10.1093/molbev/mst010",
    "21. Minh, B.Q., Schmidt, H.A., Chernomor, O., Schrempf, D., Woodhams, M.D., von Haeseler, A. and Lanfear, R. (2020) IQ-TREE 2: new models and efficient methods for phylogenetic inference in the genomic era. Mol. Biol. Evol., 37(5), 1530–1534. DOI: 10.1093/molbev/msaa015",
    "22. Kalyaanamoorthy, S., Minh, B.Q., Wong, T.K.F., von Haeseler, A. and Jermiin, L.S. (2017) ModelFinder: fast model selection for accurate phylogenetic estimates. Nat. Methods, 14(6), 587–589. DOI: 10.1038/nmeth.4285",
    "23. Guindon, S., Dufayard, J.-F., Lefort, V., Anisimova, M., Hordijk, W. and Gascuel, O. (2010) New algorithms and methods to estimate maximum-likelihood phylogenies: assessing the performance of PhyML 3.0. Syst. Biol., 59(3), 307–321. DOI: 10.1093/sysbio/syq010",
    "24. Eddy, S.R. (2011) Accelerated profile HMM searches. PLoS Comput. Biol., 7(10), e1002195. DOI: 10.1371/journal.pcbi.1002195",
    "25. Hyatt, D., Chen, G.-L., LoCascio, P.F., Land, M.L., Larimer, F.W. and Hauser, L.J. (2010) Prodigal: prokaryotic gene recognition and translation initiation site identification. BMC Bioinformatics, 11, 119. DOI: 10.1186/1471-2105-11-119",
    "26. Fu, L., Niu, B., Zhu, Z., Wu, S. and Li, W. (2012) CD-HIT: accelerated for clustering the next-generation sequencing data. Bioinformatics, 28(23), 3150–3152. DOI: 10.1093/bioinformatics/bts565",
    "27. ICTV. (2025) Virus Metadata Resource (VMR) — Master Species List #41 (MSL41). International Committee on Taxonomy of Viruses. https://ictv.global/vmr",
    "28. Lefkowitz, E.J., Dempsey, D.M., Hendrickson, R.C., Orton, R.J., Siddell, S.G. and Smith, D.B. (2018) Virus taxonomy: the database of the International Committee on Taxonomy of Viruses (ICTV). Nucleic Acids Res., 46(D1), D708–D717. DOI: 10.1093/nar/gkx932",
    "29. Arulmoorthy, M.P., Anbarasu, M., Srinivasan, M. and Thirumurugan, R. (2020) Major viral diseases in culturable penaeid shrimps: a review. Aquacult. Int., 28, 1939–1967. DOI: 10.1007/s10499-020-00568-3",
    "30. Thitamadee, S., Prachumwat, A., Srisala, J., Jaroenlak, P., Salachan, P.V., Sritunyalucksana, K., Flegel, T.W. and Itsathitphaisarn, O. (2016) Review of current disease threats for cultivated penaeid shrimp in Asia. Aquaculture, 452, 69–87. DOI: 10.1016/j.aquaculture.2015.10.028",
    "31. Lightner, D.V. (2011) Virus diseases of farmed shrimp in the Western Hemisphere (the Americas): a review. J. Invertebr. Pathol., 106(1), 110–130. DOI: 10.1016/j.jip.2010.09.012",
    "32. Dhar, A.K., Robles-Sikisaka, R., Saksmerprome, V. and Lakshman, D.K. (2014) Biology, genome organization, and evolution of parvoviruses in marine shrimp. Adv. Virus Res., 89, 85–139. DOI: 10.1016/B978-0-12-800172-1.00003-3",
    "33. Cowley, J.A., Dimmock, C.M., Wongteerasupaya, C., Boonsaeng, V., Panyim, S. and Walker, P.J. (2012) The gene 2 protein of gill-associated virus, a prawn nidovirus, is a structural protein. J. Gen. Virol., 83, 1133–1143. DOI: 10.1099/0022-1317-83-5-1133",
    "34. Naim, S., Brown, J.K. and Nibert, M.L. (2015) Genetic diversification of penaeid shrimp infectious myonecrosis virus. J. Gen. Virol., 96, 2729–2739. DOI: 10.1099/vir.0.000186",
    "35. Bonami, J.R. and Sri Widada, J. (2011) Viral diseases of the giant freshwater prawn Macrobrachium rosenbergii: a review. J. Invertebr. Pathol., 106(1), 131–142. DOI: 10.1016/j.jip.2010.09.007",
    "36. Petton, B., Destoumieux-Garzón, D., Pernet, F., Toulza, E., de Lorgeril, J., Degremont, L. and Mitta, G. (2021) The Pacific Oyster Mortality Syndrome, a polymicrobial and multifactorial disease: state of knowledge and future directions. Front. Immunol., 12, 630343. DOI: 10.3389/fimmu.2021.630343",
    "37. Whittington, R.J., Ingram, L. and Rubio, A. (2024) Environmental conditions associated with four index cases of POMS in Crassostrea gigas in Australia between 2010 and 2024. Animals, 14(21), 3052. DOI: 10.3390/ani14213052",
    "38. Burioli, E.A.V., Prearo, M. and Houssin, M. (2021) A review of the emerging disease of Pacific oyster mortality syndrome (POMS). Rev. Aquacult., 13, 1637–1653. DOI: 10.1111/raq.12541",
    "39. Coates, C.J. and Rowley, A.F. (2022) Diseases of Gastropoda. Front. Immunol., 12, 802920. DOI: 10.3389/fimmu.2021.802920",
    "40. Wang, C., Yao, L., Wang, J., Liu, G., Zhang, Q. and Li, C. (2021) First report on natural infection of nodavirus in an Echinodermata, sea cucumber (Apostichopus japonicas). Viruses, 13(4), 636. DOI: 10.3390/v13040636",
    "41. Corbeil, S., McColl, K.A., Williams, L.M., Mohammad, I., Hyatt, A.D., Crameri, S.G. and Crane, M.St.J. (2016) Abalone viral ganglioneuritis: establishment and use of an experimental immersion challenge system. J. Invertebr. Pathol., 138, 50–56. DOI: 10.1016/j.jip.2016.04.006",
    "42. Thurber, R.V., Correa, A.M.S. and Welsh, R.M. (2020) Viruses of reef-building corals. Adv. Virus Res., 107, 193–226. DOI: 10.1016/bs.aivir.2020.05.001",
    "43. Pootakham, W., Mhuantong, W., Yoocha, T., Putchim, L., Sonthirod, C., Naktang, C., Sangsrakru, D. and Tangphatsornruang, S. (2021) Taxonomy and diversity of the sponge microbiome and its associated viruses. Microbiome, 9, 1–15. DOI: 10.1186/s40168-021-01121-7",
    "44. Cárdenas, P., Pérez, T. and Boury-Esnault, N. (2020) Sponge holobiont: a multi-partner interaction model for marine systems. Front. Microbiol., 11, 536900. DOI: 10.3389/fmicb.2020.536900",
    "45. Sayers, E.W., Bolton, E.E., Brister, J.R., Canese, K., Chan, J., Comeau, D.C., Connor, R., Funk, K., Kelly, C., Kim, S., Madej, T., Marchler-Bauer, A., Lanczycki, C., Lathrop, S., Lu, Z., Thibaud-Nissen, F., Murphy, T., Phan, L., Skripchenko, Y., Tse, T., Wang, J., Williams, R., Trawick, B.W., Pruitt, K.D. and Sherry, S.T. (2022) Database resources of the National Center for Biotechnology Information. Nucleic Acids Res., 50(D1), D20–D26. DOI: 10.1093/nar/gkab1112",
    "46. The UniProt Consortium (2023) UniProt: the Universal Protein Knowledgebase in 2023. Nucleic Acids Res., 51(D1), D523–D531. DOI: 10.1093/nar/gkac1052",
    "47. Paysan-Lafosse, T., Blum, M., Chuguransky, S., Grego, T., Pinto, B.L., Salazar, G.A., Bileschi, M.L., Bork, P., Bridge, A., Colwell, L., Gough, J., Haft, D.H., Letunić, I., Marchler-Bauer, A., Mi, H., Natale, D.A., Orengo, C.A., Pandurangan, A.P., Rivoire, C., Sigrist, C.J.A., Sillitoe, I., Thanki, N., Thomas, P.D., Tosatto, S.C.E., Wu, C.H. and Bateman, A. (2023) InterPro in 2022. Nucleic Acids Res., 51(D1), D418–D427. DOI: 10.1093/nar/gkac993",
    "48. Kanehisa, M., Furumichi, M., Sato, Y., Kawashima, M. and Ishiguro-Watanabe, M. (2023) KEGG for taxonomy-based analysis of pathways and genomes. Nucleic Acids Res., 51(D1), D587–D592. DOI: 10.1093/nar/gkac963",
    "49. Mistry, J., Chuguransky, S., Williams, L., Qureshi, M., Salazar, G.A., Sonnhammer, E.L.L., Tosatto, S.C.E., Paladin, L., Raj, S., Richardson, L.J., Finn, R.D. and Bateman, A. (2021) Pfam: The protein families database in 2021. Nucleic Acids Res., 49(D1), D412–D419. DOI: 10.1093/nar/gkaa913",
    "50. Capella-Gutiérrez, S., Silla-Martínez, J.M. and Gabaldón, T. (2009) trimAl: a tool for automated alignment trimming in large-scale phylogenetic analyses. Bioinformatics, 25(15), 1972–1973. DOI: 10.1093/bioinformatics/btp348",
    "51. Minh, B.Q., Nguyen, M.A.T. and von Haeseler, A. (2013) Ultrafast approximation for phylogenetic bootstrap. Mol. Biol. Evol., 30(5), 1188–1195. DOI: 10.1093/molbev/mst024",
    "52. Kieft, K., Zhou, Z. and Anantharaman, K. (2020) VIBRANT: automated recovery, annotation and curation of microbial viruses, and evaluation of viral community function from genomic sequences. Microbiome, 8, 90. DOI: 10.1186/s40168-020-00867-0",
    "53. Nayfach, S., Camargo, A.P., Schulz, F., Eloe-Fadrosh, E., Roux, S. and Kyrpides, N.C. (2021) CheckV assesses the quality and completeness of metagenome-assembled viral genomes. Nat. Biotechnol., 39(5), 578–585. DOI: 10.1038/s41587-020-00774-7",
    "54. Rigden, D.J. and Fernández, X.M. (2025) The 2025 Nucleic Acids Research database issue and the online molecular biology database collection. Nucleic Acids Res., 53(D1), D1–D9. DOI: 10.1093/nar/gkae1251",
    "55. Wilkinson, M.D., Dumontier, M., Aalbersberg, I.J.J., Appleton, G., Axton, M., Baak, A., Blomberg, N., Boiten, J.W., da Silva Santos, L.B., Bourne, P.E., ... and Mons, B. (2016) The FAIR Guiding Principles for scientific data management and stewardship. Sci. Data, 3, 160018. DOI: 10.1038/sdata.2016.18",
    "56. GBIF.org (2024) GBIF: The Global Biodiversity Information Facility. https://www.gbif.org",
    "57. OBIS (2024) Ocean Biodiversity Information System. Intergovernmental Oceanographic Commission of UNESCO. https://obis.org",
    "58. Lin, Z., Akin, H., Rao, R., Hie, B., Zhu, Z., Lu, W., Smetanin, N., Verkuil, R., Kabeli, O., Shmueli, Y., dos Santos Costa, A., Fazel-Zarandi, M., Sercu, T., Candido, S. and Rives, A. (2023) Evolutionary-scale prediction of atomic-level protein structure with a language model. Science, 379(6637), 1123–1130. DOI: 10.1126/science.ade2574",
    "59. Callanan, J., Stockdale, S.R., Shkoporov, A., Draper, L.A., Ross, R.P. and Hill, C. (2020) Expansion of known ssRNA phage genomes: from tens to over a thousand. Sci. Adv., 6(6), eaay5981. DOI: 10.1126/sciadv.aay5981",
    "60. Tisza, M.J., Belford, A.K., Domínguez-Huerta, G., Bolduc, B. and Buck, C.B. (2021) Cenote-Taker 2 democratizes virus discovery and sequence annotation. Virus Evol., 7(1), veaa100. DOI: 10.1093/ve/veaa100",
    "61. Guo, J., Bolduc, B., Zayed, A.A., Varsani, A., Dominguez-Huerta, G., Delmont, T.O., Pratama, A.A., Gazitúa, M.C., Vik, D., Sullivan, M.B. and Marais, A. (2021) VirSorter2: a multi-classifier, expert-guided approach to detect diverse DNA and RNA viruses. Microbiome, 9, 37. DOI: 10.1186/s40168-020-00990-y",
    "62. Stentiford, G.D., Peeler, E.J., Tyler, C.R., Bickley, L.K., Holt, C.C., Bass, D., Turner, A.D., Baker-Austin, C., Ellis, T., Lowther, J.A., Posen, P.E., Bateman, K.S., Verner-Jeffreys, D.W., van Aerle, R., Stone, D.M., Paley, R., Trent, A., Katsiadaki, I., Bean, T.P. and Feist, S.W. (2020) A seafood risk tool for assessing and mitigating chemical and pathogen hazards in the aquaculture supply chain. Nat. Food, 1, 13–23. DOI: 10.1038/s43016-019-0018-2",
    "63. Walker, P.J. and Winton, J.R. (2010) Emerging viral diseases of fish and shrimp. Vet. Res., 41(6), 51. DOI: 10.1051/vetres/2010022",
    "64. Ramírez, S. (2024) FastAPI: modern, fast web framework for building APIs. https://fastapi.tiangolo.com",
    "65. Li, C.X., Shi, M., Tian, J.H., Lin, X.D., Kang, Y.J., Chen, L.J., Qin, X.C., Xu, J., Holmes, E.C. and Zhang, Y.Z. (2015) Unprecedented genomic diversity of RNA viruses in arthropods reveals the ancestry of negative-sense RNA viruses. eLife, 4, e05378. DOI: 10.7554/eLife.05378",
]

for ref in references:
    add_reference(ref)

# ============================================================
# SAVE
# ============================================================
output_path = 'F:/水生无脊椎动物数据库/NAR_PAPER_MANUSCRIPT.docx'
doc.save(output_path)
print(f'Saved: {output_path}')

import os
size_mb = os.path.getsize(output_path) / 1024 / 1024
print(f'File size: {size_mb:.1f} MB')
